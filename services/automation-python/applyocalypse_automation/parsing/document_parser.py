from __future__ import annotations

import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

from ..documents.docx_mutation import inspect_docx_for_anchors
from ..documents.tex_mutation import inspect_tex_regions

PARSER_NAME = "applyocalypse-local-parser"
PARSER_VERSION = "0.3.0"

SECTION_ALIASES = {
    "experience": {
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "employment history",
    },
    "projects": {"projects", "selected projects", "personal projects", "technical projects"},
    "education": {"education", "academic background"},
    "skills": {"skills", "technical skills", "technologies", "tools", "core skills"},
    "certifications": {"certifications", "certificates", "licenses"},
    "summary": {"summary", "profile", "professional summary"},
}

EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
PHONE_RE = re.compile(r"(?:\+?\d[\d\s().-]{7,}\d)")
URL_RE = re.compile(r"https?://[^\s)>\]]+", re.IGNORECASE)
BULLET_CHARS = "\u2022\u2023\u25aa\u25cf\u25e6\u00b7\u2043\u25a0\uf0a7\uf0b7"
BULLET_SPLIT_RE = re.compile(f"[{re.escape(BULLET_CHARS)}]")
LIST_PREFIX_RE = re.compile(rf"^\s*(?:[-*{re.escape(BULLET_CHARS)}]\s*|\d+[.)]\s+)")
SKILL_SPLIT_RE = re.compile(r"[,|;/]")
APPLYO_PLACEHOLDER_RE = re.compile(r"\{\{APPLYO_[A-Z0-9_]+\}\}")
SHORT_HEADING_MAX_CHARS = 120

# Word packs several visual lines into a single paragraph using soft breaks and
# inline bullet glyphs, and uses tab runs as column gutters. Resumes carry the
# company / date / location triple in those columns, so the layout has to be
# recovered before any section or heading logic can see what a human sees.
COLUMN_SEPARATOR = " | "
SOFT_BREAK_RE = re.compile(r"[\r\n\v\f]+")
TAB_RUN_RE = re.compile(r"\t+")
COLUMN_RUN_RE = re.compile(r"(?:\s*\|\s*)+")

_MONTH = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?"
_DATE_TOKEN = rf"(?:{_MONTH}\s+\d{{4}}|{_MONTH}\s*'\d{{2}}|\d{{1,2}}/\d{{4}}|\d{{4}})"
_OPEN_ENDED = r"(?:Present|Current|Ongoing|Now)"
_DASH = "\\-\u2010\u2011\u2012\u2013\u2014\u2015\u2212~"
DATE_RANGE_RE = re.compile(
    rf"(?P<start>{_DATE_TOKEN})\s*(?:[{_DASH}]|to|through|until)\s*(?P<end>{_DATE_TOKEN}|{_OPEN_ENDED})",
    re.IGNORECASE,
)
DATE_ONLY_RE = re.compile(rf"(?:{_DATE_TOKEN}|{_OPEN_ENDED})", re.IGNORECASE)
LOCATION_RE = re.compile(r"[A-Z][\w.'-]*(?:[ .][A-Z][\w.'-]*)*,\s*(?:[A-Z]{2}|[A-Z][a-z]+(?: [A-Z][a-z]+)?)")
DEGREE_RE = re.compile(
    r"\b(?:bachelors?|masters?|associates?|doctorate|doctoral|ph\.?\s?d|b\.?\s?s\.?|b\.?\s?a\.?"
    r"|m\.?\s?s\.?|m\.?\s?a\.?|b\.?\s?tech|m\.?\s?tech|b\.?\s?e\.?|m\.?\s?eng|mba|diploma)\b",
    re.IGNORECASE,
)
HEADING_DATE_REMAINDER_MAX_WORDS = 6
CERT_YEAR_RE = re.compile(r"\(\s*(?P<year>(?:19|20)\d{2})\s*\)")


@dataclass(frozen=True, slots=True)
class ParsedSection:
    section_id: str
    label: str
    normalized_label: str
    start_line: int
    end_line: int
    confidence: float
    items: list[str] = field(default_factory=list)

    def to_payload(self) -> dict[str, Any]:
        return {
            "sectionId": self.section_id,
            "label": self.label,
            "normalizedLabel": self.normalized_label,
            "startLine": self.start_line,
            "endLine": self.end_line,
            "confidence": self.confidence,
            "items": self.items,
        }


@dataclass(frozen=True, slots=True)
class ParsedDocumentResult:
    parser_name: str
    parser_version: str
    confidence: float
    canonical: dict[str, Any]
    style_map: dict[str, Any]
    anchor_map: dict[str, Any]
    warnings: list[str]

    def to_payload(self) -> dict[str, Any]:
        return {
            "parser_name": self.parser_name,
            "parser_version": self.parser_version,
            "confidence": self.confidence,
            "canonical": self.canonical,
            "style_map": self.style_map,
            "anchor_map": self.anchor_map,
            "warnings": self.warnings,
        }


def _source_format(path: Path) -> str:
    extension = path.suffix.lower()
    if extension == ".pdf":
        return "PDF"
    if extension == ".docx":
        return "DOCX"
    if extension == ".tex":
        return "TEX"
    if extension == ".md":
        return "MD"
    if extension == ".txt":
        return "TXT"
    raise ValueError(f"Unsupported document format: {extension or '(none)'}")


def _document_kind(value: str) -> str:
    normalized = value.upper()
    if normalized in {"RESUME", "COVER_LETTER", "SUPPORTING_DETAILS", "OTHER"}:
        return normalized
    return "OTHER"


def _normalize_line(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip())


def _tidy_columns(value: str) -> str:
    collapsed = COLUMN_RUN_RE.sub(COLUMN_SEPARATOR, value)
    return _normalize_line(collapsed).strip("|").strip()


def _expand_layout_lines(value: str) -> list[str]:
    """Recover the visual lines packed into one paragraph or table cell."""
    lines: list[str] = []
    for soft_line in SOFT_BREAK_RE.split(value):
        for chunk in BULLET_SPLIT_RE.split(soft_line):
            line = _tidy_columns(TAB_RUN_RE.sub(COLUMN_SEPARATOR, chunk))
            if line:
                lines.append(line)
    return lines


def _layout_lines(text: str) -> list[str]:
    lines: list[str] = []
    for raw in text.splitlines():
        lines.extend(_expand_layout_lines(raw))
    return lines


def _run_segments(paragraph: Any) -> list[tuple[str, bool]]:
    """(text, italic) per run, including runs nested inside hyperlink elements.

    Word writes hyperlinks as a `w:hyperlink` wrapper whose inner `w:r` holds the
    visible text. python-docx reads only direct `w:r` children, so certification
    names and linked project titles vanish from `Paragraph.text`.
    """
    from docx.oxml.ns import qn  # type: ignore

    segments: list[tuple[str, bool]] = []
    for run in paragraph._p.iter(qn("w:r")):  # noqa: SLF001
        text = ""
        for node in run:
            if node.tag == qn("w:t"):
                text += node.text or ""
            elif node.tag == qn("w:tab"):
                text += "\t"
            elif node.tag in {qn("w:br"), qn("w:cr")}:
                text += "\n"
        if not text:
            continue  # wrapper run: its hyperlink child carries the text
        run_properties = run.find(qn("w:rPr"))
        italic_node = None if run_properties is None else run_properties.find(qn("w:i"))
        italic = italic_node is not None and italic_node.get(qn("w:val")) not in {"0", "false"}
        segments.append((text, italic))
    return segments


def _paragraph_layout_text(paragraph: Any) -> str:
    segments = _run_segments(paragraph)
    text = "".join(segment for segment, _ in segments)
    if "\t" in text or len(text) > SHORT_HEADING_MAX_CHARS:
        return text

    # Resume templates mark the role in italics where another template would use
    # a tab stop, so treat the switch into italics as a column gutter.
    parts: list[str] = []
    previous_italic = False
    for index, (segment, italic) in enumerate(segments):
        if italic and not previous_italic and index > 0 and "".join(parts).strip():
            parts.append("\t")
        previous_italic = italic
        parts.append(segment)
    return "".join(parts)


def _cell_layout_lines(cell: Any) -> list[str]:
    lines: list[str] = []
    for paragraph in cell.paragraphs:
        lines.extend(_expand_layout_lines(_paragraph_layout_text(paragraph)))
    for nested in cell.tables:
        lines.extend(_table_layout_lines(nested))
    return lines


def _table_layout_lines(table: Any) -> list[str]:
    """Two-column resume tables read row by row, pairing each cell's Nth line."""
    lines: list[str] = []
    for row in table.rows:
        columns: list[list[str]] = []
        seen: set[Any] = set()
        for cell in row.cells:
            if cell._tc in seen:  # noqa: SLF001 - merged cells repeat in row.cells
                continue
            seen.add(cell._tc)  # noqa: SLF001
            columns.append(_cell_layout_lines(cell))
        depth = max((len(column) for column in columns), default=0)
        for offset in range(depth):
            parts = [column[offset] for column in columns if offset < len(column)]
            line = _tidy_columns(COLUMN_SEPARATOR.join(parts))
            if line:
                lines.append(line)
    return lines


def _plain_text_from_docx(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from docx import Document  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.table import Table  # type: ignore
        from docx.text.paragraph import Paragraph  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX parsing") from exc

    document = Document(str(path))
    paragraphs: list[dict[str, Any]] = []
    text_lines: list[str] = []
    paragraph_index = 0
    table_index = 0

    # Body order matters: python-docx exposes paragraphs and tables as separate
    # collections, and flattening them that way moves every table to the end of
    # the document, so section attribution lands on the wrong heading.
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            expanded = _expand_layout_lines(_paragraph_layout_text(paragraph))
            if expanded:
                style = paragraph.style.name if paragraph.style is not None else ""
                is_bullet = paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None  # noqa: SLF001
                paragraphs.append(
                    {
                        "index": paragraph_index,
                        "style": style,
                        "isBullet": is_bullet,
                        "runCount": len(paragraph.runs),
                        "textPreview": expanded[0][:160],
                    }
                )
                text_lines.extend(expanded)
            paragraph_index += 1
        elif child.tag == qn("w:tbl"):
            for row_offset, line in enumerate(_table_layout_lines(Table(child, document))):
                paragraphs.append(
                    {
                        "index": f"table:{table_index}:{row_offset}",
                        "style": "table-row",
                        "isBullet": False,
                        "runCount": 0,
                        "textPreview": line[:160],
                    }
                )
                text_lines.append(line)
            table_index += 1

    return "\n".join(text_lines), {"paragraphs": paragraphs, "tableCount": len(document.tables)}


def _plain_text(path: Path) -> tuple[str, dict[str, Any], list[str]]:
    source_format = _source_format(path)
    warnings: list[str] = []
    if source_format == "DOCX":
        text, style_map = _plain_text_from_docx(path)
        return text, style_map, warnings
    if source_format in {"TEX", "TXT", "MD"}:
        return path.read_text(encoding="utf-8", errors="replace"), {}, warnings
    warnings.append("PDF is ingestion-only. Parse the verified DOCX editable master produced from conversion.")
    return "", {}, warnings


def _detect_applyo_placeholders(source: str) -> list[str]:
    placeholders: list[str] = []
    seen: set[str] = set()
    for match in APPLYO_PLACEHOLDER_RE.finditer(source):
        placeholder = match.group(0)
        if placeholder not in seen:
            seen.add(placeholder)
            placeholders.append(placeholder)
    return placeholders


def _canonical_section_label(line: str) -> tuple[str, float] | None:
    raw = _normalize_line(line).strip(":")
    cleaned = re.sub(r"[^a-z0-9 &/+-]", "", raw.lower())
    if len(cleaned) > 42:
        return None
    for canonical, aliases in SECTION_ALIASES.items():
        if cleaned in aliases:
            return canonical, 0.86
    # Test the pre-lowercased text: an ALL-CAPS heading marks a section boundary
    # even when it is not a known alias (e.g. AWARDS, PUBLICATIONS, LANGUAGES).
    if raw.isupper() and 2 <= len(raw) <= 32 and cleaned:
        return cleaned, 0.58
    return None


def _strip_list_prefix(value: str) -> str:
    return LIST_PREFIX_RE.sub("", value).strip()


def _extract_sections(lines: list[str]) -> list[ParsedSection]:
    markers: list[tuple[int, str, str, float]] = []
    for index, line in enumerate(lines):
        label = _canonical_section_label(line)
        if label:
            markers.append((index, line, label[0], label[1]))

    sections: list[ParsedSection] = []
    for marker_index, (line_index, label, normalized_label, confidence) in enumerate(markers):
        end_line = markers[marker_index + 1][0] if marker_index + 1 < len(markers) else len(lines)
        items = [
            _strip_list_prefix(line)
            for line in lines[line_index + 1 : end_line]
            if _strip_list_prefix(line) and not _canonical_section_label(line)
        ]
        sections.append(
            ParsedSection(
                section_id=f"section:{normalized_label}:{marker_index}",
                label=label,
                normalized_label=normalized_label,
                start_line=line_index,
                end_line=max(line_index, end_line - 1),
                confidence=confidence,
                items=items,
            )
        )
    return sections


def _identity_from_lines(lines: list[str]) -> dict[str, Any]:
    joined = "\n".join(lines[:18])
    email = EMAIL_RE.search(joined)
    phone = PHONE_RE.search(joined)
    links = [
        {"label": url.split("//", 1)[-1].split("/", 1)[0], "url": url.rstrip(".,")}
        for url in URL_RE.findall(joined)
    ]

    legal_name: str | None = None
    for line in lines[:6]:
        if EMAIL_RE.search(line) or URL_RE.search(line) or PHONE_RE.search(line):
            continue
        candidate = _normalize_line(line)
        if 2 <= len(candidate.split()) <= 5 and len(candidate) <= 80:
            legal_name = candidate
            break

    return {
        "legalName": legal_name,
        "email": email.group(0) if email else None,
        "phone": _normalize_line(phone.group(0)) if phone else None,
        "location": None,
        "links": links,
    }


def _skills_from_section(section: ParsedSection) -> list[str]:
    skills: list[str] = []
    for item in section.items:
        cleaned = re.sub(r"^[A-Za-z /+&-]{2,32}:\s*", "", item)
        parts = [part.strip() for part in SKILL_SPLIT_RE.split(cleaned)]
        for part in parts:
            if not part or len(part) > 48:
                continue
            if len(part.split()) > 5:
                continue
            skills.append(part)
    deduped: list[str] = []
    seen: set[str] = set()
    for skill in skills:
        key = skill.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(skill)
    return deduped


def _split_once(value: str, separators: list[str]) -> tuple[str, str] | None:
    for separator in separators:
        if separator in value:
            left, right = value.split(separator, 1)
            left = _normalize_line(left)
            right = _normalize_line(right)
            if left and right:
                return left, right
    return None


def _is_probable_heading(value: str) -> bool:
    item = _normalize_line(value)
    if not item or len(item) > SHORT_HEADING_MAX_CHARS:
        return False
    if item.endswith((".", ";")):
        return False
    return True


def _split_columns(item: str) -> list[str]:
    return [segment for segment in (part.strip() for part in item.split("|")) if segment]


def _heading_fields(
    *,
    company: str | None,
    title: str | None,
    location: str | None = None,
    start_date: str | None = None,
    end_date: str | None = None,
) -> dict[str, Any]:
    return {
        "company": company,
        "title": title,
        "location": location,
        "startDate": start_date,
        "endDate": end_date,
    }


def _date_range(segment: str) -> tuple[str, str, str] | None:
    """Pull a date range out of one column, returning the leftover column text."""
    match = DATE_RANGE_RE.search(segment)
    if match is None:
        return None
    remainder = _normalize_line(f"{segment[: match.start()]} {segment[match.end() :]}")
    if len(remainder.split()) > HEADING_DATE_REMAINDER_MAX_WORDS:
        return None  # prose that merely mentions dates, not an entry heading
    return _normalize_line(match.group("start")), _normalize_line(match.group("end")), remainder


def _dated_heading(item: str) -> dict[str, Any] | None:
    """A resume row carrying a date range is an entry heading; the other columns
    are the company, title and location in whatever order the template used."""
    dates: tuple[str, str] | None = None
    columns: list[str] = []
    for segment in _split_columns(item):
        parsed = _date_range(segment) if dates is None else None
        if parsed is None:
            columns.append(segment)
            continue
        start, end, remainder = parsed
        dates = (start, end)
        if remainder:
            columns.append(remainder)
    if dates is None:
        return None
    location = next((segment for segment in columns[1:] if LOCATION_RE.fullmatch(segment)), None)
    named = [segment for segment in columns if segment != location]
    return _heading_fields(
        company=named[0] if named else None,
        title=named[1] if len(named) > 1 else None,
        location=location,
        start_date=dates[0],
        end_date=dates[1],
    )


def _entry_detail_only(item: str) -> dict[str, str] | None:
    """A row carrying only a location or only dates continues the entry above it."""
    fields: dict[str, str] = {}
    for segment in _split_columns(item):
        parsed = _date_range(segment)
        if parsed is not None:
            start, end, remainder = parsed
            if remainder or "startDate" in fields:
                return None
            fields["startDate"], fields["endDate"] = start, end
        elif DATE_ONLY_RE.fullmatch(segment):
            fields.setdefault("endDate", segment)
        elif LOCATION_RE.fullmatch(segment):
            fields.setdefault("location", segment)
        else:
            return None
    return fields or None


def _experience_heading(item: str) -> dict[str, Any] | None:
    if not _is_probable_heading(item):
        return None
    dated = _dated_heading(item)
    if dated and dated["company"]:
        return dated
    at_match = re.match(r"^(?P<title>.+?)\s+(?:at|@)\s+(?P<company>.+)$", item, flags=re.IGNORECASE)
    if at_match:
        return _heading_fields(
            company=_normalize_line(at_match.group("company")),
            title=_normalize_line(at_match.group("title")),
        )

    split = _split_once(item, [" | ", " - "])
    if split:
        left, right = split
        if len(left.split()) <= 8 and len(right.split()) <= 10:
            return _heading_fields(company=right, title=left)
    return None


def _experience_from_section(section: ParsedSection) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for item in section.items:
        heading = _experience_heading(item)
        if heading:
            if current:
                entries.append(current)
            current = {
                **heading,
                "bullets": [],
                "tools": [],
                "confidence": min(section.confidence, 0.82),
            }
            continue
        if not current or not item:
            continue
        detail = _entry_detail_only(item)
        if detail:
            for key, value in detail.items():
                current[key] = current[key] or value
            continue
        current["bullets"].append(item)
    if current:
        entries.append(current)
    return entries


def _project_from_items(section: ParsedSection) -> list[dict[str, Any]]:
    projects: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for item in section.items:
        if _is_probable_heading(item):
            split = _split_once(item, [": ", " - ", " | "])
            if split or current is None:
                if current:
                    projects.append(current)
                name, summary = split if split else (item, None)
                current = {
                    "name": name,
                    "role": None,
                    "summary": summary,
                    "bullets": [],
                    "tools": [],
                    "links": URL_RE.findall(item),
                    "confidence": min(section.confidence, 0.82),
                }
                continue
        if current and item:
            current["bullets"].append(item)
    if current:
        projects.append(current)
    return projects


def _split_institution_and_degree(value: str) -> tuple[str, str | None]:
    """Templates often run the institution and the degree together on one line."""
    match = DEGREE_RE.search(value)
    if match is None or match.start() == 0:
        return value, None
    institution = _normalize_line(value[: match.start()]).strip("|,-").strip()
    degree = _normalize_line(value[match.start() :])
    if not institution or not degree:
        return value, None
    return institution, degree


def _education_columns(item: str) -> tuple[list[str], str | None, str | None, str | None]:
    """Classify one education row into named columns, dates and a location."""
    named: list[str] = []
    start_date: str | None = None
    end_date: str | None = None
    location: str | None = None
    for segment in _split_columns(item):
        parsed = _date_range(segment) if end_date is None else None
        if parsed is not None:
            start_date, end_date, remainder = parsed
            if remainder:
                named.append(remainder)
        elif DATE_ONLY_RE.fullmatch(segment):
            end_date = end_date or segment
        elif location is None and LOCATION_RE.fullmatch(segment):
            location = segment
        else:
            named.append(segment)
    return named, start_date, end_date, location


def _education_from_section(section: ParsedSection) -> list[dict[str, Any]]:
    education: list[dict[str, Any]] = []
    for item in section.items:
        if not _is_probable_heading(item):
            continue
        named, start_date, end_date, location = _education_columns(item)
        details = [location] if location else []

        previous = education[-1] if education else None
        # A row with no institution, or a degree-only row, continues the entry
        # above it: two-column templates split one school across two lines.
        continues = previous is not None and (
            not named or (previous["degree"] is None and DEGREE_RE.match(named[0]) is not None)
        )
        if continues and previous is not None:
            if named:
                previous["degree"] = named[0]
                details.extend(named[1:])
            previous["startDate"] = previous["startDate"] or start_date
            previous["endDate"] = previous["endDate"] or end_date
            previous["details"].extend(details)
            continue
        if not named:
            continue

        institution, degree = _split_institution_and_degree(named[0])
        if degree is None and len(named) > 1:
            degree = named[1]
            details.extend(named[2:])
        else:
            details.extend(named[1:])
        education.append(
            {
                "institution": institution,
                "degree": degree,
                "field": None,
                "startDate": start_date,
                "endDate": end_date,
                "details": details,
                "confidence": min(section.confidence, 0.8),
            }
        )
    return education


def _certification_from_item(item: str, confidence: float) -> dict[str, Any]:
    """Split a certification line into its name and the year it was issued.

    Resume templates commonly append "(2025)" plus a skill list to the
    certification name; keeping that tail would store it as part of the name.
    """
    match = CERT_YEAR_RE.search(item)
    name = _normalize_line(item[: match.start()]).rstrip(",;:-") if match else item
    return {
        "name": name.strip() or item,
        "issuer": None,
        "issuedAt": match.group("year") if match else None,
        "expiresAt": None,
        "credentialUrl": None,
        "confidence": confidence,
    }


def _canonical_from_sections(*, source_format: str, document_kind: str, text: str, sections: list[ParsedSection]) -> dict[str, Any]:
    lines = _layout_lines(text)
    skill_groups = []
    for section in sections:
        if section.normalized_label == "skills":
            skills = _skills_from_section(section)
            if skills:
                skill_groups.append({"label": section.label, "skills": skills, "confidence": min(section.confidence, 0.82)})

    education = [
        entry
        for section in sections
        if section.normalized_label == "education"
        for entry in _education_from_section(section)
    ]
    experience = [
        entry
        for section in sections
        if section.normalized_label == "experience"
        for entry in _experience_from_section(section)
    ]
    projects = [
        entry
        for section in sections
        if section.normalized_label == "projects"
        for entry in _project_from_items(section)
    ]
    certifications = [
        _certification_from_item(item, min(section.confidence, 0.78))
        for section in sections
        if section.normalized_label == "certifications"
        for item in section.items
        if item
    ]

    return {
        "documentKind": document_kind,
        "sourceFormat": source_format,
        "identity": _identity_from_lines(lines),
        "sections": [section.to_payload() for section in sections],
        "skillGroups": skill_groups,
        "education": education,
        "experience": experience,
        "projects": projects,
        "certifications": certifications,
        "rawTextPreview": text if document_kind == "COVER_LETTER" else text[:2000],
    }


def parse_document(path: Path, *, document_kind: str = "OTHER") -> ParsedDocumentResult:
    source_format = _source_format(path)
    text, style_map, warnings = _plain_text(path)
    lines = _layout_lines(text)
    sections = _extract_sections(lines)
    anchor_map: dict[str, Any] = {}

    if source_format == "DOCX":
        plan = inspect_docx_for_anchors(path)
        placeholders = _detect_applyo_placeholders(text)
        anchor_map = {"anchors": [asdict(anchor) for anchor in plan.anchors], "placeholders": placeholders}
        warnings.extend(plan.warnings)
        if not placeholders:
            warnings.append("No explicit Applyocalypse placeholders found. Automated DOCX mutation will pause for anchor repair.")
    elif source_format == "TEX":
        regions = inspect_tex_regions(path)
        source = path.read_text(encoding="utf-8", errors="replace")
        placeholders = _detect_applyo_placeholders(source)
        anchor_map = {"regions": [asdict(region) for region in regions], "placeholders": placeholders}
        if not regions:
            warnings.append("No high-confidence TEX regions found. User review is required before automated mutation.")
        if not placeholders:
            warnings.append("No explicit Applyocalypse placeholders found. Automated TEX mutation will pause for anchor repair.")

    if source_format == "PDF":
        confidence = 0.15
    else:
        section_confidence = sum(section.confidence for section in sections) / len(sections) if sections else 0.45
        anchor_bonus = 0.08 if anchor_map else 0
        confidence = min(0.92, max(0.35, section_confidence + anchor_bonus))

    if not sections and source_format != "PDF":
        warnings.append("No recognizable resume sections were detected. Manual profile review is required.")

    canonical = _canonical_from_sections(
        source_format=source_format,
        document_kind=_document_kind(document_kind),
        text=text,
        sections=sections,
    )

    return ParsedDocumentResult(
        parser_name=PARSER_NAME,
        parser_version=PARSER_VERSION,
        confidence=round(confidence, 3),
        canonical=canonical,
        style_map=style_map,
        anchor_map=anchor_map,
        warnings=warnings,
    )
