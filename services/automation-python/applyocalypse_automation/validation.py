from __future__ import annotations

import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Literal

BANNED_WORDS = [
    "leverage",
    "utilize",
    "spearheaded",
    "robust",
    "comprehensive",
    "seamless",
    "transformative",
    "passionate",
    "excited",
    "thrilled",
    "dynamic",
    "innovative",
    "holistic",
    "empower",
    "foster",
    "harness",
    "deep dive",
]

WEAK_BULLET_PATTERNS = [
    r"\bresponsible for\b",
    r"\bworked on\b",
    r"\bhelped with\b",
    r"\bparticipated in\b",
    r"\bvarious\b",
]


@dataclass(frozen=True, slots=True)
class ValidationIssue:
    code: str
    message: str
    severity: str


@dataclass(frozen=True, slots=True)
class ValidationReport:
    passed: bool
    blocking_issues: list[ValidationIssue]
    warnings: list[ValidationIssue]

    def human_report(self) -> str:
        lines = ["Validation passed." if self.passed else "Validation failed."]
        if self.blocking_issues:
            lines.append("Blocking issues:")
            lines.extend(f"- {issue.code}: {issue.message}" for issue in self.blocking_issues)
        if self.warnings:
            lines.append("Warnings:")
            lines.extend(f"- {issue.code}: {issue.message}" for issue in self.warnings)
        return "\n".join(lines)

    def to_dict(self) -> dict[str, object]:
        return {
            "passed": self.passed,
            "blocking_issues": [asdict(issue) for issue in self.blocking_issues],
            "warnings": [asdict(issue) for issue in self.warnings],
            "human_report": self.human_report(),
        }


ArtifactKind = Literal["resume", "cover_letter", "generic"]


@dataclass(frozen=True, slots=True)
class ArtifactTextExtraction:
    text: str
    source_format: str
    warnings: list[str]


def extract_docx_text(path: Path) -> ArtifactTextExtraction:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX validation") from exc

    document = Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = " ".join(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
                if cell_text:
                    lines.append(cell_text)

    return ArtifactTextExtraction(text="\n".join(lines), source_format="DOCX", warnings=[])


def extract_tex_text(path: Path) -> ArtifactTextExtraction:
    source = path.read_text(encoding="utf-8")
    warnings: list[str] = []

    try:
        from TexSoup import TexSoup  # type: ignore
    except ImportError:
        warnings.append("TexSoup unavailable; used constrained TEX text fallback.")
        return ArtifactTextExtraction(text=extract_tex_text_fallback(source), source_format="TEX", warnings=warnings)

    try:
        soup = TexSoup(source)
        text = str(soup.text)
        if text.strip():
            return ArtifactTextExtraction(text=normalize_extracted_text(text), source_format="TEX", warnings=warnings)
    except Exception as exc:
        warnings.append(f"TexSoup parsing failed; used constrained fallback: {type(exc).__name__}.")

    return ArtifactTextExtraction(text=extract_tex_text_fallback(source), source_format="TEX", warnings=warnings)


def extract_tex_text_fallback(source: str) -> str:
    without_comments = re.sub(r"(?<!\\)%.*", "", source)
    without_environments = re.sub(r"\\(begin|end)\s*\{[^}]+\}", "\n", without_comments)
    with_item_markers = re.sub(r"\\item\b", "\n- ", without_environments)
    without_commands_with_args = re.sub(r"\\[a-zA-Z*]+\s*(?:\[[^\]]*])?\s*\{([^{}]*)\}", r"\1", with_item_markers)
    without_commands = re.sub(r"\\[a-zA-Z*]+(?:\s*\[[^\]]*])?", " ", without_commands_with_args)
    unescaped = (
        without_commands.replace(r"\&", "&")
        .replace(r"\%", "%")
        .replace(r"\_", "_")
        .replace(r"\#", "#")
        .replace(r"\$", "$")
    )
    return normalize_extracted_text(unescaped)


def extract_pdf_text(path: Path) -> ArtifactTextExtraction:
    try:
        import fitz  # type: ignore
    except ImportError as exc:
        raise RuntimeError("PyMuPDF is required for rendered PDF validation") from exc

    lines: list[str] = []
    with fitz.open(path) as document:
        for page in document:
            text = page.get_text("text").strip()
            if text:
                lines.append(text)

    return ArtifactTextExtraction(text="\n".join(lines), source_format="PDF", warnings=[])


def normalize_extracted_text(text: str) -> str:
    normalized_lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in normalized_lines if line)


def extract_artifact_text(path: Path) -> ArtifactTextExtraction:
    if not path.exists():
        raise FileNotFoundError(path)

    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return ArtifactTextExtraction(text=path.read_text(encoding="utf-8"), source_format=suffix[1:].upper(), warnings=[])
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".tex":
        return extract_tex_text(path)
    if suffix == ".pdf":
        return extract_pdf_text(path)
    raise ValueError(f"Unsupported validation artifact format: {suffix or '<none>'}")


_BULLET_CHAR_LIMIT_DEFAULT = 240

# Em dash plus visually-equivalent variants; en dash (U+2013) stays legal for date ranges.
_EM_DASH_VARIANTS = "—―⸺⸻﹘"


class TextArtifactValidator:
    def validate(
        self,
        text: str,
        *,
        artifact_kind: ArtifactKind,
        bullet_char_limit: int = _BULLET_CHAR_LIMIT_DEFAULT,
    ) -> ValidationReport:
        blocking: list[ValidationIssue] = []
        warnings: list[ValidationIssue] = []
        lower = text.lower()

        for word in BANNED_WORDS:
            pattern = r"\b" + re.escape(word).replace(r"\ ", r"\s+") + r"\b"
            if re.search(pattern, lower):
                blocking.append(ValidationIssue("BANNED_WORD", f"Banned wording found: {word}", "blocking"))

        em_dash_count = sum(text.count(char) for char in _EM_DASH_VARIANTS)
        if em_dash_count:
            blocking.append(ValidationIssue("EM_DASH", f"Em dash count must be 0. Found {em_dash_count}.", "blocking"))

        if artifact_kind == "cover_letter":
            word_count = len([token for token in re.split(r"\s+", text.strip()) if token])
            if word_count > 450:
                blocking.append(ValidationIssue("COVER_LETTER_TOO_LONG", f"Cover letter has {word_count} words.", "blocking"))
            if re.match(r"^\s*(i am applying for|i'?m writing to apply|i am writing to apply)\b", text, flags=re.IGNORECASE):
                warnings.append(ValidationIssue("GENERIC_OPENER", "Cover letter starts with a generic opener.", "warning"))
            if not re.search(r"\b(sincerely|best regards|regards|thank you)\b", text, flags=re.IGNORECASE):
                warnings.append(ValidationIssue("MISSING_SIGN_OFF", "Cover letter should include a human sign-off.", "warning"))

        if artifact_kind == "resume":
            bullets = [line for line in text.splitlines() if re.match(r"^\s*[-*]\s+", line)]
            for bullet in bullets:
                word_count = len([token for token in re.split(r"\s+", bullet.strip()) if token])
                if word_count > 32:
                    warnings.append(ValidationIssue("LONG_BULLET", "Resume bullet is longer than the target length.", "warning"))
                bullet_text = re.sub(r"^\s*[-*]\s+", "", bullet)
                if len(bullet_text) > bullet_char_limit:
                    warnings.append(ValidationIssue(
                        "BULLET_TOO_LONG",
                        f"Resume bullet exceeds {bullet_char_limit} characters ({len(bullet_text)} chars).",
                        "warning",
                    ))
                if any(re.search(pattern, bullet, flags=re.IGNORECASE) for pattern in WEAK_BULLET_PATTERNS):
                    warnings.append(ValidationIssue("WEAK_BULLET", "Resume bullet uses weak or vague phrasing.", "warning"))

            for required in [r"\bexperience\b", r"\bskills?\b"]:
                if not re.search(required, text, flags=re.IGNORECASE):
                    warnings.append(ValidationIssue("RESUME_MISSING_SECTION", "Resume may be missing an expected section.", "warning"))

            first_words = [
                re.sub(r"^\s*[-*]\s+", "", bullet).strip().split()[0].lower()
                for bullet in bullets
                if re.sub(r"^\s*[-*]\s+", "", bullet).strip().split()
            ]
            if any(first_words.count(word) >= 3 for word in set(first_words)):
                warnings.append(ValidationIssue("REPETITIVE_RHYTHM", "Multiple bullets start with the same word.", "warning"))

        return ValidationReport(passed=not blocking, blocking_issues=blocking, warnings=warnings)

    def validate_file(
        self,
        path: Path,
        *,
        artifact_kind: ArtifactKind,
        bullet_char_limit: int = _BULLET_CHAR_LIMIT_DEFAULT,
    ) -> ValidationReport:
        extraction = extract_artifact_text(path)
        report = self.validate(extraction.text, artifact_kind=artifact_kind, bullet_char_limit=bullet_char_limit)
        extraction_warnings = [
            ValidationIssue("TEXT_EXTRACTION_WARNING", warning, "warning") for warning in extraction.warnings
        ]
        length_warnings: list[ValidationIssue] = []
        if artifact_kind == "resume" and path.suffix.lower() in {".docx", ".tex"}:
            length_warnings = self.check_docx_page_count(path)
        return ValidationReport(
            passed=report.passed,
            blocking_issues=report.blocking_issues,
            warnings=[*extraction_warnings, *length_warnings, *report.warnings],
        )

    @staticmethod
    def check_docx_page_count(docx_path: Path) -> list[ValidationIssue]:
        """One-page check: uses real PDF page count when a companion PDF exists,
        otherwise falls back to a heuristic character-count estimate.

        Companion PDF is expected at the same path with a .pdf extension.
        Flags RESUME_LENGTH_WARNING if page count > 1.
        """
        from applyocalypse_automation.documents.font_detection import count_pdf_pages

        companion_pdf = docx_path.with_suffix(".pdf")
        if companion_pdf.exists():
            pages = count_pdf_pages(companion_pdf)
            if pages > 1:
                return [ValidationIssue(
                    "RESUME_LENGTH_WARNING",
                    f"Resume is {pages} pages. Trim to one page for best ATS results.",
                    "warning",
                )]
            if pages > 0:
                return []

        # Heuristic fallback: total non-whitespace chars / 3500
        try:
            from docx import Document  # type: ignore[import]
            doc = Document(str(docx_path))
            total_chars = sum(len(p.text.strip()) for p in doc.paragraphs if p.text.strip())
            estimated_pages = total_chars / 3500
            if estimated_pages > 1.2:
                return [ValidationIssue(
                    "RESUME_LENGTH_WARNING",
                    f"Resume may exceed one page (estimated {estimated_pages:.1f} pages). "
                    "Trim to keep ATS scanners reading all content.",
                    "warning",
                )]
        except Exception:
            pass
        return []
