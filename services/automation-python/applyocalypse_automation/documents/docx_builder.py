from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from typing import Any


def _string_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def build_cover_letter_docx(
    text: str,
    canonical_profile: dict[str, Any],
    output_path: Path,
) -> None:
    """Write a plain-text cover letter as a DOCX using python-docx.

    Format: Calibri 11, 1-inch normal margins, name + contact header,
    today's date, then body paragraphs (split on double newline).
    """
    try:
        from docx import Document  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError("python-docx is required for cover letter DOCX generation") from exc

    profile = canonical_profile.get("profile") if isinstance(canonical_profile.get("profile"), dict) else {}
    legal_name = str(profile.get("legalName") or profile.get("displayName") or "").strip()
    email = str(profile.get("email") or "").strip()
    phone = str(profile.get("phone") or "").strip()
    location = str(profile.get("location") or "").strip()
    contact_parts = [part for part in [email, phone, location] if part]

    doc = Document()

    # Margins: 1 inch on all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def _set_run_font(run: Any, bold: bool = False) -> None:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.bold = bold
        # Force theme font override so Word renders Calibri
        run._r.get_or_add_rPr()
        rFonts = run._r.rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")

    def _add_para(text_content: str, bold: bool = False) -> Any:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text_content)
        _set_run_font(run, bold=bold)
        return p

    # Header: name line
    if legal_name:
        _add_para(legal_name, bold=True)
    if contact_parts:
        _add_para("  |  ".join(contact_parts))

    # Blank line between header and date
    _add_para("")

    # Date
    today = datetime.now(UTC).strftime("%B %d, %Y")
    _add_para(today)

    # Blank line before body
    _add_para("")

    # Body: split on blank lines; each block is a paragraph
    body_blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
    for i, block in enumerate(body_blocks):
        # Within each block, replace single newlines with spaces for paragraph continuity
        single_line = " ".join(block.splitlines())
        _add_para(single_line)
        if i < len(body_blocks) - 1:
            _add_para("")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def build_resume_docx(
    canonical_profile: dict[str, Any],
    tailoring_plan: dict[str, Any],
    output_path: Path,
    *,
    font_size: int = 10,
) -> None:
    """Build a plain-text resume DOCX from canonical profile and tailoring plan.

    Used as the anchor-free fallback when the DOCX master has no mutation anchors.
    Format: Calibri {font_size}pt, 0.75-inch margins, name/contact header, then
    Skills, Experience (bullets), Projects, and Education sections.
    """
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Inches, Pt, RGBColor  # type: ignore  # noqa: F401
    except ImportError as exc:
        raise RuntimeError("python-docx is required for resume DOCX generation") from exc

    profile = canonical_profile.get("profile") if isinstance(canonical_profile.get("profile"), dict) else {}
    legal_name = str(profile.get("legalName") or profile.get("displayName") or "").strip()
    email = str(profile.get("email") or "").strip()
    phone = str(profile.get("phone") or "").strip()
    location = str(profile.get("location") or "").strip()
    linkedin = str(profile.get("linkedin") or "").strip()
    contact_parts = [part for part in [email, phone, location, linkedin] if part]

    one_page_plan = tailoring_plan.get("one_page_plan") if isinstance(tailoring_plan.get("one_page_plan"), dict) else {}
    bullet_limit = int(one_page_plan.get("bullet_limit_per_role", 4)) if isinstance(one_page_plan.get("bullet_limit_per_role"), (int, float)) else 4
    experience_limit = int(one_page_plan.get("experience_limit", 4)) if isinstance(one_page_plan.get("experience_limit"), (int, float)) else 4
    project_limit = int(one_page_plan.get("project_limit", 3)) if isinstance(one_page_plan.get("project_limit"), (int, float)) else 3

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    pt = Pt(font_size)

    def _set_run(run: Any, bold: bool = False, size: Any = None) -> None:
        run.font.name = "Calibri"
        run.font.size = size or pt
        run.bold = bold
        run._r.get_or_add_rPr()
        rFonts = run._r.rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")

    def _add_plain(text_content: str, bold: bool = False, center: bool = False, size: Any = None) -> Any:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text_content)
        _set_run(run, bold=bold, size=size)
        return p

    def _add_section_heading(text_content: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text_content.upper())
        _set_run(run, bold=True)
        # Horizontal rule below heading via bottom border
        from docx.oxml import OxmlElement  # type: ignore
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_bullet(text_content: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text_content)
        _set_run(run)

    # Name header
    if legal_name:
        _add_plain(legal_name, bold=True, center=True, size=Pt(font_size + 4))
    if contact_parts:
        _add_plain("  |  ".join(contact_parts), center=True)

    # Skills
    skill_groups = canonical_profile.get("skillGroups") if isinstance(canonical_profile.get("skillGroups"), list) else []
    all_skills: list[str] = []
    for grp in skill_groups:
        if isinstance(grp, dict):
            all_skills.extend(_string_list(grp.get("skills")))
    skills_priority = _string_list(one_page_plan.get("skills_priority"))
    ordered_skills: list[str] = [
        s for priority in skills_priority for s in all_skills if s.lower() == priority.lower()
    ]
    ordered_skills.extend([s for s in all_skills if s not in ordered_skills])
    if ordered_skills:
        _add_section_heading("Skills")
        _add_plain(", ".join(ordered_skills[:36]))

    # Experience
    experience = canonical_profile.get("experience") if isinstance(canonical_profile.get("experience"), list) else []
    if experience:
        _add_section_heading("Experience")
        for entry in experience[:experience_limit]:
            if not isinstance(entry, dict):
                continue
            title = str(entry.get("title") or "").strip()
            company_name = str(entry.get("company") or "").strip()
            heading = " | ".join(part for part in [title, company_name] if part)
            if heading:
                _add_plain(heading, bold=True)
            bullets_raw = _string_list(entry.get("bullets"))
            for bullet in bullets_raw[:bullet_limit]:
                _add_bullet(bullet)

    # Projects
    projects = canonical_profile.get("projects") if isinstance(canonical_profile.get("projects"), list) else []
    if projects:
        _add_section_heading("Projects")
        for project in projects[:project_limit]:
            if not isinstance(project, dict):
                continue
            name = str(project.get("name") or "").strip()
            if name:
                _add_plain(name, bold=True)
            summary = str(project.get("summary") or "").strip()
            if summary:
                _add_plain(summary)
            for bullet in _string_list(project.get("bullets"))[:bullet_limit]:
                _add_bullet(bullet)

    # Education
    education = canonical_profile.get("education") if isinstance(canonical_profile.get("education"), list) else []
    if education:
        _add_section_heading("Education")
        for entry in education[:3]:
            if not isinstance(entry, dict):
                continue
            institution = str(entry.get("institution") or "").strip()
            degree = str(entry.get("degree") or "").strip()
            field_name = str(entry.get("field") or "").strip()
            line = " | ".join(part for part in [institution, degree, field_name] if part)
            if line:
                _add_plain(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
