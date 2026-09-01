from __future__ import annotations

import copy
from collections.abc import Iterable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass(frozen=True, slots=True)
class DocxAnchor:
    anchor_id: str
    section_label: str
    paragraph_index: int
    run_fingerprint: str
    confidence: float
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True, slots=True)
class DocxMutationPlan:
    source_docx: Path
    anchors: list[DocxAnchor]
    warnings: list[str]


@dataclass(frozen=True, slots=True)
class DocxParagraphMutation:
    paragraph_index: int
    replacement_text: str


def iter_table_paragraphs(table: Any) -> Iterable[Any]:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def iter_document_paragraphs(document: Any) -> list[Any]:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        paragraphs.extend(iter_table_paragraphs(table))
    return paragraphs


def addressable_paragraphs(document: Any) -> list[Any]:
    """Every paragraph a mutation can target, in a stable order: body paragraphs first,
    then table cells depth-first, each appearing exactly once.

    Two reasons this is not just ``iter_document_paragraphs``. A merged cell is returned
    once per grid column it spans, so the same paragraph comes back several times, and an
    index has to name exactly one paragraph for a mutation to be well defined. Identity is
    the underlying XML element rather than the wrapper, because python-docx builds a fresh
    ``Paragraph`` on every access.
    """
    seen: set[int] = set()
    unique: list[Any] = []
    for paragraph in iter_document_paragraphs(document):
        key = id(paragraph._element)
        if key in seen:
            continue
        seen.add(key)
        unique.append(paragraph)
    return unique


def extract_docx_text(path: Path) -> str:
    """Return plain text from all paragraphs and table cells in a DOCX file."""
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX text extraction") from exc
    document = Document(str(path))
    return "\n".join(p.text for p in iter_document_paragraphs(document) if p.text.strip())


def replace_paragraph_text_preserving_runs(paragraph: Any, replacement_text: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(replacement_text)
        return

    original_lengths = [max(len(run.text), 0) for run in runs]
    total_original_length = sum(original_lengths)
    if total_original_length == 0:
        runs[0].text = replacement_text
        for run in runs[1:]:
            run.text = ""
        return

    cursor = 0
    replacement_length = len(replacement_text)
    for index, run in enumerate(runs):
        if index == len(runs) - 1:
            run.text = replacement_text[cursor:]
            break

        share = original_lengths[index] / total_original_length
        next_cursor = min(replacement_length, cursor + round(replacement_length * share))
        run.text = replacement_text[cursor:next_cursor]
        cursor = next_cursor


def mutate_docx_paragraphs(source_docx: Path, output_docx: Path, mutations: list[DocxParagraphMutation]) -> Path:
    if source_docx.suffix.lower() != ".docx":
        raise ValueError("source_docx must be a DOCX file")
    if output_docx.suffix.lower() != ".docx":
        raise ValueError("output_docx must be a DOCX file")
    if not source_docx.exists():
        raise FileNotFoundError(source_docx)

    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX mutation") from exc

    document = Document(str(source_docx))
    paragraphs = addressable_paragraphs(document)
    for mutation in mutations:
        if mutation.paragraph_index < 0 or mutation.paragraph_index >= len(paragraphs):
            raise IndexError(f"paragraph index out of range: {mutation.paragraph_index}")
        replace_paragraph_text_preserving_runs(paragraphs[mutation.paragraph_index], mutation.replacement_text)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))
    return output_docx


_BULLET_GLYPHS = "•◦‣▪·-–*"


def collect_tailorable_bullets(document: Any) -> list[tuple[int, str]]:
    """Paragraphs that look like substantive resume bullets (list style or a bullet
    glyph), returned as (paragraph_index, text). Indices are stable against a fresh
    re-open of the same file, so they can drive ``mutate_docx_paragraphs`` directly.

    Table cells count. Plenty of resume templates are one invisible table, or two columns,
    and reading only the body would find no bullets at all in those and tailor nothing
    while still reporting success."""
    bullets: list[tuple[int, str]] = []
    for index, paragraph in enumerate(addressable_paragraphs(document)):
        text = paragraph.text.strip()
        if len(text) < 25:
            continue
        try:
            style_name = (paragraph.style.name or "").lower()
        except Exception:  # noqa: BLE001 - a missing style must not abort detection
            style_name = ""
        if "list" in style_name or text[:1] in _BULLET_GLYPHS:
            bullets.append((index, text))
    return bullets


def parses_back_intact(
    before: list[str], after: list[str], tailored_by_index: dict[int, str]
) -> bool:
    """Whether a written document still reads back as the document we meant to write:
    the same paragraphs, the tailored ones holding exactly their new text, the rest
    untouched.

    Rewriting a paragraph means redistributing text across its existing runs, which is how
    the formatting survives, and it is also how a character goes missing at a run boundary
    or a space disappears from a ``w:t`` that lost ``xml:space``. An ATS parses the file
    rather than rendering it, so a resume that still looks right in Word can already be
    wrong by the time it is read."""
    if len(before) != len(after):
        return False
    for index, (original, written) in enumerate(zip(before, after, strict=True)):
        expected = tailored_by_index.get(index, original)
        if written != expected:
            return False
    return True


def tailor_master_docx_in_place(
    master_docx: Path, output_docx: Path, tailored_by_index: dict[int, str]
) -> tuple[Path, int]:
    """Write tailored bullet text back into the master DOCX in place, preserving every
    paragraph's font/size/bullet/indent (only run text changes). Copies the master
    verbatim when there is nothing to change, so the user's exact layout is always kept.

    The written file is read back and checked against what was asked for. If it does not
    match, the master is copied through untouched and the count comes back 0: an
    untailored resume the user recognises beats a tailored one with a mangled line."""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX mutation") from exc

    def copy_master_verbatim() -> tuple[Path, int]:
        output_docx.parent.mkdir(parents=True, exist_ok=True)
        Document(str(master_docx)).save(str(output_docx))
        return output_docx, 0

    mutations = [
        DocxParagraphMutation(paragraph_index=index, replacement_text=text)
        for index, text in sorted(tailored_by_index.items())
    ]
    if not mutations:
        return copy_master_verbatim()

    before = [p.text for p in addressable_paragraphs(Document(str(master_docx)))]
    mutate_docx_paragraphs(master_docx, output_docx, mutations)
    after = [p.text for p in addressable_paragraphs(Document(str(output_docx)))]
    if not parses_back_intact(before, after, tailored_by_index):
        return copy_master_verbatim()
    return output_docx, len(mutations)


def mutate_docx_placeholders(source_docx: Path, output_docx: Path, replacements: dict[str, str]) -> tuple[Path, list[str]]:
    if source_docx.suffix.lower() != ".docx":
        raise ValueError("source_docx must be a DOCX file")
    if output_docx.suffix.lower() != ".docx":
        raise ValueError("output_docx must be a DOCX file")
    if not source_docx.exists():
        raise FileNotFoundError(source_docx)

    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX mutation") from exc

    document = Document(str(source_docx))
    replaced: list[str] = []
    for paragraph in iter_document_paragraphs(document):
        paragraph_text = paragraph.text
        next_text = paragraph_text
        for placeholder, replacement in replacements.items():
            if placeholder in next_text:
                next_text = next_text.replace(placeholder, replacement)
                replaced.append(placeholder)
        if next_text != paragraph_text:
            replace_paragraph_text_preserving_runs(paragraph, next_text)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))
    return output_docx, sorted(set(replaced))


def _expand_anchor_to_bullets(document: Any, anchor_text: str, bullets: list[str]) -> bool:
    """Find the paragraph containing anchor_text and replace it with one paragraph per bullet.

    New paragraphs copy the original paragraph's XML (style, list numbering, indentation)
    so bullet formatting is preserved exactly. Returns True if the anchor was found.
    """
    try:
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required for bullet expansion") from exc

    if not bullets:
        return False

    anchor_para = None
    for para in iter_document_paragraphs(document):
        if anchor_text in para.text:
            anchor_para = para
            break

    if anchor_para is None:
        return False

    p_elem = anchor_para._p  # noqa: SLF001
    parent = p_elem.getparent()
    insert_idx = list(parent).index(p_elem)

    orig_runs = p_elem.findall(qn("w:r"))
    orig_rPr = orig_runs[0].find(qn("w:rPr")) if orig_runs else None

    for offset, bullet_text in enumerate(bullets):
        new_p = copy.deepcopy(p_elem)
        for r in new_p.findall(qn("w:r")):
            new_p.remove(r)
        r_elem = OxmlElement("w:r")
        if orig_rPr is not None:
            r_elem.append(copy.deepcopy(orig_rPr))
        t_elem = OxmlElement("w:t")
        t_elem.text = bullet_text
        if bullet_text.startswith(" ") or bullet_text.endswith(" "):
            t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_elem.append(t_elem)
        new_p.append(r_elem)
        parent.insert(insert_idx + offset, new_p)

    parent.remove(p_elem)
    return True


def mutate_docx_bullet_anchors(
    source_docx: Path,
    output_docx: Path,
    bullet_replacements: dict[str, list[str]],
) -> tuple[Path, list[str]]:
    """Replace bullet anchor placeholders with expanded multi-paragraph bullet lists.

    Each anchor paragraph (e.g. '{{APPLYO_EXP_0_BULLETS}}') is expanded into one
    paragraph per bullet string, preserving the original paragraph style and list formatting.
    """
    if source_docx.suffix.lower() != ".docx":
        raise ValueError("source_docx must be a DOCX file")
    if not source_docx.exists():
        raise FileNotFoundError(source_docx)

    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX mutation") from exc

    document = Document(str(source_docx))
    replaced: list[str] = []
    for anchor, bullets in bullet_replacements.items():
        if _expand_anchor_to_bullets(document, anchor, bullets):
            replaced.append(anchor)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))
    return output_docx, replaced


def inspect_docx_for_anchors(source_docx: Path) -> DocxMutationPlan:
    if source_docx.suffix.lower() != ".docx":
        raise ValueError("source_docx must be a DOCX file")
    if not source_docx.exists():
        raise FileNotFoundError(source_docx)

    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX inspection") from exc

    document = Document(str(source_docx))
    anchors: list[DocxAnchor] = []
    warnings: list[str] = []

    for index, paragraph in enumerate(iter_document_paragraphs(document)):
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        is_likely_heading = "heading" in style_name.lower()
        is_likely_bullet = paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None  # noqa: SLF001
        if is_likely_heading or is_likely_bullet:
            fingerprint = "|".join(run.text[:24] for run in paragraph.runs if run.text)
            anchors.append(
                DocxAnchor(
                    anchor_id=f"docx:p:{index}",
                    section_label=text[:80],
                    paragraph_index=index,
                    run_fingerprint=fingerprint,
                    confidence=0.72 if is_likely_bullet else 0.66,
                    metadata={"style": style_name, "bullet": is_likely_bullet},
                )
            )

    if not anchors:
        warnings.append("No high-confidence anchors found. User review is required before automated mutation.")

    return DocxMutationPlan(source_docx=source_docx, anchors=anchors, warnings=warnings)
