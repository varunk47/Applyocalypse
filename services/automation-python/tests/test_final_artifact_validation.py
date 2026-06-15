from __future__ import annotations

from pathlib import Path

import pytest

try:
    from docx import Document
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

from applyocalypse_automation.validation import TextArtifactValidator


def test_validate_file_blocks_em_dash_in_docx(tmp_path: Path) -> None:
    if not _DOCX_AVAILABLE:
        pytest.skip("python-docx not installed")

    docx_path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("Led teams — delivered results")
    doc.save(str(docx_path))

    report = TextArtifactValidator().validate_file(docx_path, artifact_kind="resume")

    assert report.passed is False
    codes = [issue.code for issue in report.blocking_issues]
    assert any("EM_DASH" in code for code in codes), f"Expected EM_DASH in blocking codes, got: {codes}"


def test_validate_file_passes_clean_docx(tmp_path: Path) -> None:
    if not _DOCX_AVAILABLE:
        pytest.skip("python-docx not installed")

    docx_path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("Led teams and delivered results across multiple products")
    doc.save(str(docx_path))

    report = TextArtifactValidator().validate_file(docx_path, artifact_kind="resume")

    assert report.passed is True


def test_bullet_retry_message_includes_violation_codes() -> None:
    from applyocalypse_automation.runner import _build_bullet_retry_jd

    # Build a failing ValidationReport by validating a string with an em dash
    failing_text = "Led teams — delivered results"
    report = TextArtifactValidator().validate(failing_text, artifact_kind="resume")
    assert not report.passed, "Expected validation to fail for em-dash input"

    job_text = "Software Engineer at Acme Corp"
    result = _build_bullet_retry_jd(job_text, report)

    assert job_text in result
    assert "EM_DASH" in result
    assert "No em dashes" in result
