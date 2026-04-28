from __future__ import annotations

from pathlib import Path

from docx import Document

from applyocalypse_automation.documents.anchor_repair import repair_editable_master_anchors
from applyocalypse_automation.documents.docx_mutation import DocxParagraphMutation, mutate_docx_paragraphs, mutate_docx_placeholders
from applyocalypse_automation.documents.file_generation import (
    GeneratedNameInput,
    build_generated_filename,
    choose_collision_safe_path,
)
from applyocalypse_automation.documents.pdf_export import export_docx_to_pdf
from applyocalypse_automation.documents.tex_mutation import (
    TexBlockMutation,
    compile_tex_with_tectonic,
    inspect_tex_regions,
    mutate_tex_placeholders,
    mutate_tex_regions,
)
from applyocalypse_automation.parsing.document_parser import parse_document


def test_docx_mutation_preserves_run_structure(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "tailored.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Built ")
    paragraph.add_run("systems")
    document.save(source)

    mutate_docx_paragraphs(
        source,
        output,
        [DocxParagraphMutation(paragraph_index=0, replacement_text="Built distributed systems")],
    )

    tailored = Document(output)
    assert tailored.paragraphs[0].text == "Built distributed systems"
    assert len(tailored.paragraphs[0].runs) == 2


def test_docx_placeholder_mutation_is_explicit_and_run_aware(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "tailored.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Summary: ")
    paragraph.add_run("{{APPLYO_RESUME_SUMMARY}}")
    document.save(source)

    _, replaced = mutate_docx_placeholders(source, output, {"{{APPLYO_RESUME_SUMMARY}}": "Verified Python evidence."})

    tailored = Document(output)
    assert replaced == ["{{APPLYO_RESUME_SUMMARY}}"]
    assert tailored.paragraphs[0].text == "Summary: Verified Python evidence."
    assert len(tailored.paragraphs[0].runs) == 2


def test_docx_placeholder_mutation_reaches_table_cells(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "tailored.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.add_run("Skills: ")
    paragraph.add_run("{{APPLYO_SKILLS}}")
    document.save(source)

    _, replaced = mutate_docx_placeholders(source, output, {"{{APPLYO_SKILLS}}": "Python, SQLite"})

    tailored = Document(output)
    cell_paragraph = tailored.tables[0].cell(0, 0).paragraphs[0]
    assert replaced == ["{{APPLYO_SKILLS}}"]
    assert cell_paragraph.text == "Skills: Python, SQLite"
    assert len(cell_paragraph.runs) == 2


def test_parser_reports_docx_applyocalypse_placeholders(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Summary: {{APPLYO_RESUME_SUMMARY}}")
    document.save(source)

    parsed = parse_document(source, document_kind="RESUME")

    assert parsed.anchor_map["placeholders"] == ["{{APPLYO_RESUME_SUMMARY}}"]
    assert not any("No explicit Applyocalypse placeholders" in warning for warning in parsed.warnings)


def test_parser_detects_docx_sections_inside_table_cells(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    document = Document()
    table = document.add_table(rows=4, cols=1)
    table.cell(0, 0).text = "Summary"
    table.cell(1, 0).text = "Built deterministic engines."
    table.cell(2, 0).text = "Skills"
    table.cell(3, 0).text = "Python, SQLite"
    document.save(source)

    parsed = parse_document(source, document_kind="RESUME")
    sections = {section["normalizedLabel"]: section for section in parsed.canonical["sections"]}

    assert "summary" in sections
    assert "skills" in sections
    assert parsed.canonical["skillGroups"][0]["skills"] == ["Python", "SQLite"]
    assert parsed.style_map["tableCount"] == 1


def test_parser_extracts_structured_profile_sections_from_text(tmp_path: Path) -> None:
    source = tmp_path / "resume.txt"
    source.write_text(
        "\n".join(
            [
                "Grace Hopper",
                "grace@example.com",
                "Experience",
                "Principal Engineer at Compiler Systems Lab",
                "Built deterministic developer tooling with Python and SQLite",
                "Projects",
                "Application Control Plane: Local-first automation review console",
                "Designed approval gates for sensitive browser actions",
                "Education",
                "Yale University - PhD Mathematics",
                "Certifications",
                "Naval Computing Systems",
                "Skills",
                "Python, TypeScript, SQLite",
            ]
        ),
        encoding="utf-8",
    )

    parsed = parse_document(source, document_kind="RESUME")
    canonical = parsed.canonical

    assert canonical["identity"]["email"] == "grace@example.com"
    assert canonical["experience"][0]["company"] == "Compiler Systems Lab"
    assert canonical["experience"][0]["title"] == "Principal Engineer"
    assert canonical["experience"][0]["bullets"] == ["Built deterministic developer tooling with Python and SQLite"]
    assert canonical["projects"][0]["name"] == "Application Control Plane"
    assert canonical["education"][0]["institution"] == "Yale University"
    assert canonical["certifications"][0]["name"] == "Naval Computing Systems"
    assert canonical["skillGroups"][0]["skills"] == ["Python", "TypeScript", "SQLite"]


def test_docx_anchor_repair_creates_reviewable_candidate_without_touching_source(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "anchored.docx"
    document = Document()
    document.add_paragraph("Ada Lovelace")
    document.add_paragraph("Summary")
    document.add_paragraph("Built deterministic engines.")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, SQLite")
    document.save(source)

    result = repair_editable_master_anchors(source, output)
    source_document = Document(source)
    repaired_document = Document(output)

    assert result.added_placeholders == ["{{APPLYO_FULL_NAME}}", "{{APPLYO_RESUME_SUMMARY}}", "{{APPLYO_SKILLS}}"]
    assert source_document.paragraphs[0].text == "Ada Lovelace"
    assert repaired_document.paragraphs[0].text == "{{APPLYO_FULL_NAME}}"
    assert "{{APPLYO_SKILLS}}" in "\n".join(paragraph.text for paragraph in repaired_document.paragraphs)


def test_docx_anchor_repair_detects_section_bodies_inside_tables(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "anchored.docx"
    document = Document()
    document.add_paragraph("Ada Lovelace")
    table = document.add_table(rows=4, cols=1)
    table.cell(0, 0).text = "Summary"
    table.cell(1, 0).text = "Built deterministic engines."
    table.cell(2, 0).text = "Skills"
    table.cell(3, 0).text = "Python, SQLite"
    document.save(source)

    result = repair_editable_master_anchors(source, output)
    repaired_document = Document(output)
    table_text = "\n".join(
        paragraph.text for row in repaired_document.tables[0].rows for cell in row.cells for paragraph in cell.paragraphs
    )

    assert result.added_placeholders == ["{{APPLYO_FULL_NAME}}", "{{APPLYO_RESUME_SUMMARY}}", "{{APPLYO_SKILLS}}"]
    assert "{{APPLYO_RESUME_SUMMARY}}" in table_text
    assert "{{APPLYO_SKILLS}}" in table_text


def test_tex_fallback_mutation_replaces_target_region(tmp_path: Path) -> None:
    source = tmp_path / "resume.tex"
    output = tmp_path / "tailored.tex"
    source.write_text("\\section{Experience}\n\\begin{itemize}\n\\item Built APIs\n\\end{itemize}\n", encoding="utf-8")

    regions = inspect_tex_regions(source)
    item_region = next(region for region in regions if region.command_name == "item")
    mutate_tex_regions(source, output, [TexBlockMutation(region_id=item_region.region_id, replacement_source="\\item Built reliable APIs")])

    output_text = output.read_text(encoding="utf-8")
    assert "\\item Built reliable APIs" in output_text
    assert "\\section{Experience}" in output_text


def test_tex_placeholder_mutation_preserves_source_structure(tmp_path: Path) -> None:
    source = tmp_path / "resume.tex"
    output = tmp_path / "tailored.tex"
    source.write_text("\\section{Skills}\n{{APPLYO_SKILLS}}\n", encoding="utf-8")

    _, replaced = mutate_tex_placeholders(source, output, {"{{APPLYO_SKILLS}}": "Python, SQLite"})

    assert replaced == ["{{APPLYO_SKILLS}}"]
    assert output.read_text(encoding="utf-8") == "\\section{Skills}\nPython, SQLite\n"


def test_parser_warns_when_tex_lacks_applyocalypse_placeholders(tmp_path: Path) -> None:
    source = tmp_path / "resume.tex"
    source.write_text("\\section{Skills}\nPython, SQLite\n", encoding="utf-8")

    parsed = parse_document(source, document_kind="RESUME")

    assert parsed.anchor_map["placeholders"] == []
    assert any("No explicit Applyocalypse placeholders" in warning for warning in parsed.warnings)


def test_tex_anchor_repair_replaces_section_bodies_without_touching_source(tmp_path: Path) -> None:
    source = tmp_path / "resume.tex"
    output = tmp_path / "anchored.tex"
    source.write_text(
        "\\documentclass{article}\n\\begin{document}\n\\section{Summary}\nBuilt tools.\n\\section{Skills}\nPython, SQLite\n\\end{document}\n",
        encoding="utf-8",
    )

    result = repair_editable_master_anchors(source, output)
    repaired = output.read_text(encoding="utf-8")

    assert result.added_placeholders == ["{{APPLYO_RESUME_SUMMARY}}", "{{APPLYO_SKILLS}}"]
    assert "Built tools." in source.read_text(encoding="utf-8")
    assert "{{APPLYO_RESUME_SUMMARY}}" in repaired
    assert "{{APPLYO_SKILLS}}" in repaired


def test_tex_compile_reports_missing_tectonic(monkeypatch, tmp_path: Path) -> None:
    source = tmp_path / "resume.tex"
    source.write_text("\\documentclass{article}\\begin{document}Hi\\end{document}", encoding="utf-8")

    def raise_missing(*_args, **_kwargs):
        raise FileNotFoundError("tectonic")

    monkeypatch.setattr("subprocess.run", raise_missing)

    result = compile_tex_with_tectonic(source, tmp_path / "out")

    assert result.ok is False
    assert result.pdf_path is None
    assert "Tectonic executable was not found" in result.stderr


def test_docx_pdf_export_reports_missing_local_exporter(monkeypatch, tmp_path: Path) -> None:
    source = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Resume")
    document.save(source)

    monkeypatch.setattr("shutil.which", lambda _command: None)

    result = export_docx_to_pdf(source, tmp_path / "out")

    assert result.ok is False
    assert result.pdf_path is None
    assert result.code == "DOCX_PDF_EXPORTER_UNAVAILABLE"


def test_generated_filename_and_collision_policy(tmp_path: Path) -> None:
    filename = build_generated_filename(
        GeneratedNameInput(
            first_name="Ada",
            last_name="Lovelace",
            company="Analytical/Engines",
            role="Staff Engineer",
            kind="Resume",
            extension="PDF",
        )
    )
    first = choose_collision_safe_path(tmp_path, filename)
    first.write_text("existing", encoding="utf-8")
    second = choose_collision_safe_path(tmp_path, filename)

    assert filename == "Ada Lovelace Analytical Engines Staff Engineer Resume.pdf"
    assert second.name == "Ada Lovelace Analytical Engines Staff Engineer Resume (2).pdf"
