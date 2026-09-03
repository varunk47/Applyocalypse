from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from applyocalypse_automation.documents import ats_lint
from applyocalypse_automation.documents.ats_lint import lint_resume_docx

# A real 1x1 PNG, so add_picture produces the a:blip a genuine Word document has.
_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)

_VML_TEXT_BOX = (
    '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:v="urn:schemas-microsoft-com:vml">'
    "<w:r><w:pict><v:shape><v:textbox><w:txbxContent>"
    "<w:p><w:r><w:t>Jane Doe, jane@example.com</w:t></w:r></w:p>"
    "</w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
)

# A clean single-column resume with the headings Greenhouse recognises. Every
# fixture starts from this so a test only ever adds the one hazard it is about.
_CLEAN_LINES = (
    "Jane Doe",
    "jane@example.com | 555-0100 | Austin, TX",
    "Experience",
    "Senior Platform Engineer, Acme Inc. | January 2022 to Present",
    "Cut deploy time from forty minutes to six by parallelising the build.",
    "Education",
    "Bachelor of Science, Computer Science",
    "Skills",
    "Python, TypeScript, PostgreSQL, Kubernetes",
)


def _write(tmp_path: Path, document: Document, name: str = "resume.docx") -> Path:
    path = tmp_path / name
    document.save(str(path))
    return path


def _clean_document(lines: tuple[str, ...] = _CLEAN_LINES) -> Document:
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    return document


def _codes(path: Path) -> set[str]:
    return {risk.code for risk in lint_resume_docx(path)}


def test_a_clean_single_column_resume_reports_nothing(tmp_path: Path) -> None:
    """The baseline every other test is measured against: no false alarms."""
    assert _codes(_write(tmp_path, _clean_document())) == set()


def test_tables_are_reported(tmp_path: Path) -> None:
    document = _clean_document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Senior Platform Engineer"
    table.rows[0].cells[1].text = "January 2022 to Present"
    assert _codes(_write(tmp_path, document)) == {"TABLES"}


def test_a_text_box_is_reported_once_not_twice(tmp_path: Path) -> None:
    """Word nests w:txbxContent inside v:textbox, so a naive sum would say two."""
    document = _clean_document()
    document.element.body.append(parse_xml(_VML_TEXT_BOX))
    path = _write(tmp_path, document)

    risks = [risk for risk in lint_resume_docx(path) if risk.code == "TEXT_BOXES"]
    assert len(risks) == 1
    assert "1 text box" in risks[0].detail


def test_an_embedded_image_is_reported(tmp_path: Path) -> None:
    document = _clean_document()
    document.add_picture(BytesIO(_PNG_1X1))
    assert _codes(_write(tmp_path, document)) == {"GRAPHICS"}


def test_a_multi_column_page_is_reported(tmp_path: Path) -> None:
    document = _clean_document()
    document.sections[0]._sectPr.find(qn("w:cols")).set(qn("w:num"), "2")
    assert _codes(_write(tmp_path, document)) == {"MULTI_COLUMN"}


@pytest.mark.parametrize(("where", "code"), [("header", "HEADER_CONTENT"), ("footer", "FOOTER_CONTENT")])
def test_content_in_the_page_furniture_is_reported(tmp_path: Path, where: str, code: str) -> None:
    document = _clean_document()
    getattr(document.sections[0], where).paragraphs[0].text = "Jane Doe | jane@example.com"
    assert _codes(_write(tmp_path, document)) == {code}


@pytest.mark.parametrize("where", ["header", "footer"])
def test_an_empty_header_or_footer_is_not_reported(tmp_path: Path, where: str) -> None:
    """Word writes the part whether or not anything is in it, so presence proves nothing."""
    document = _clean_document()
    getattr(document.sections[0], where).paragraphs[0].text = "   "
    assert _codes(_write(tmp_path, document)) == set()


def test_letters_spaced_apart_are_reported(tmp_path: Path) -> None:
    document = _clean_document((*_CLEAN_LINES, "P R O J E C T S"))
    codes = _codes(_write(tmp_path, document))
    assert codes == {"SPACED_LETTERS"}


@pytest.mark.parametrize(
    ("label", "line"),
    [
        ("ordinary_prose", "Built a b testing infrastructure for the growth team."),
        ("initials", "J. D. Rockefeller Fellowship, 2019"),
        ("single_letters_far_apart", "Grade A in one course and B in another across two terms."),
    ],
)
def test_ordinary_text_does_not_trip_the_spaced_letter_check(tmp_path: Path, label: str, line: str) -> None:
    """Four single characters in a row is the threshold precisely to avoid these."""
    document = _clean_document((*_CLEAN_LINES, line))
    assert "SPACED_LETTERS" not in _codes(_write(tmp_path, document))


def test_a_resume_with_invented_section_names_is_reported(tmp_path: Path) -> None:
    document = _clean_document(
        (
            "Jane Doe",
            "jane@example.com",
            "Where I Have Made An Impact",
            "Senior Platform Engineer, Acme Inc. | January 2022 to Present",
            "What I Studied",
            "Bachelor of Science, Computer Science",
        )
    )
    assert _codes(_write(tmp_path, document)) == {"NO_STANDARD_HEADINGS"}


@pytest.mark.parametrize("heading", ["Experience", "WORK HISTORY", "Technical Skills:", "  Education  "])
def test_one_recognised_heading_is_enough(tmp_path: Path, heading: str) -> None:
    """Case, a trailing colon and surrounding whitespace must not change the answer."""
    document = _clean_document(("Jane Doe", "jane@example.com", heading, "Acme Inc. | January 2022 to Present"))
    assert "NO_STANDARD_HEADINGS" not in _codes(_write(tmp_path, document))


def test_a_document_with_no_text_layer_reports_that_and_stops(tmp_path: Path) -> None:
    """Otherwise the missing-headings check fires too and says the same thing twice."""
    document = Document()
    document.add_picture(BytesIO(_PNG_1X1))
    assert _codes(_write(tmp_path, document)) == {"GRAPHICS", "NO_TEXT_LAYER"}


def test_an_oversized_file_is_reported(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """The threshold is patched rather than the file inflated: this tests the branch,
    not the filesystem's ability to hold 2.5MB."""
    monkeypatch.setattr(ats_lint, "_MAX_PARSEABLE_BYTES", 100)
    risks = [risk for risk in lint_resume_docx(_write(tmp_path, _clean_document())) if risk.code == "FILE_TOO_LARGE"]
    assert len(risks) == 1
    assert "2.5MB" in risks[0].detail


def test_every_risk_serialises_for_an_event_payload(tmp_path: Path) -> None:
    document = _clean_document()
    document.add_table(rows=1, cols=2)
    payloads = [risk.to_payload() for risk in lint_resume_docx(_write(tmp_path, document))]
    assert payloads == [{"code": "TABLES", "detail": payloads[0]["detail"]}]
    assert payloads[0]["detail"]
