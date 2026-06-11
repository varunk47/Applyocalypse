from __future__ import annotations

from pathlib import Path

from applyocalypse_automation.answers import propose_answer_for_detected_field
from applyocalypse_automation.documents.artifact_generation import (
    build_cover_letter_text,
    build_resume_markdown,
    split_legal_name,
    write_text_artifact,
)
from applyocalypse_automation.documents.file_generation import GeneratedNameInput
from applyocalypse_automation.jd_analysis import JobDescriptionAnalyzer, normalize_llm_analysis
from applyocalypse_automation.parsing.canonical_profile import CanonicalProfileBuilder, MergeSource, ProfileFact
from applyocalypse_automation.tailoring.engine import TailoringEngine
from applyocalypse_automation.validation import TextArtifactValidator


def test_jd_analysis_extracts_tools_and_cover_letter_signal() -> None:
    analysis = JobDescriptionAnalyzer().analyze(
        "Senior engineer required: Python, TypeScript, and SQLite. Please include a cover letter."
    )

    assert "python" in analysis.tools_and_technologies
    assert analysis.cover_letter_likely_required is True
    assert "cover_letter" in analysis.required_materials


def test_llm_analysis_normalization_preserves_fallback_for_bad_fields() -> None:
    fallback = JobDescriptionAnalyzer().analyze("Required: Python. Cover letter required.")
    normalized = normalize_llm_analysis(
        {
            "must_have_keywords": ["Python", "SQLite"],
            "preferred_keywords": "bad-shape",
            "cover_letter_likely_required": True,
        },
        fallback,
    )

    assert normalized.must_have_keywords == ["Python", "SQLite"]
    assert normalized.preferred_keywords == fallback.preferred_keywords
    assert normalized.cover_letter_likely_required is True


def test_canonical_profile_user_edit_wins() -> None:
    builder = CanonicalProfileBuilder(
        facts=[
            ProfileFact("email", "parsed@example.com", MergeSource.PARSED_TEXT, 0.95),
            ProfileFact("email", "user@example.com", MergeSource.USER_EDIT, 0.8),
        ]
    )

    canonical = builder.merge()

    assert canonical["email"]["value"] == "user@example.com"
    assert canonical["email"]["source"] == "USER_EDIT"


def test_tailoring_engine_ranks_verified_evidence_without_adding_missing_keywords() -> None:
    canonical_profile = {
        "experience": [
            {
                "id": "exp-1",
                "title": "Principal Engineer",
                "company": "Analytical Engines",
                "bullets": [
                    "Built Python automation for SQLite workflows.",
                    "Reviewed TypeScript services for deterministic release checks.",
                ],
                "tools": ["Python", "SQLite", "TypeScript"],
                "confidence": 0.94,
            }
        ],
        "projects": [
            {
                "id": "project-1",
                "name": "Portal Control Center",
                "summary": "Electron and SolidJS monitoring surface for automation review.",
                "bullets": ["Created a SolidJS console for field review and run diagnostics."],
                "tools": ["Electron", "SolidJS"],
                "confidence": 0.88,
            }
        ],
        "skillGroups": [{"id": "skills-1", "label": "Skills", "skills": ["Python", "SQLite", "SolidJS"], "confidence": 1.0}],
        "education": [],
    }
    jd_analysis = {
        "must_have_keywords": ["Python", "SQLite", "Kubernetes"],
        "preferred_keywords": ["SolidJS"],
        "tools_and_technologies": ["Electron"],
        "cover_letter_likely_required": True,
        "risky_claims_to_avoid": ["Kubernetes production ownership"],
    }

    plan = TailoringEngine().build_plan(canonical_profile=canonical_profile, jd_analysis=jd_analysis).to_dict()

    matched_keywords = {item["keyword"] for item in plan["matched_evidence"]}
    assert {"Python", "SQLite", "SolidJS", "Electron"}.issubset(matched_keywords)
    assert "Kubernetes" in plan["missing_keywords"]
    assert all("Kubernetes" not in " ".join(item.get("keep_bullets", [])) for item in plan["resume_bullet_plan"])
    assert plan["resume_bullet_plan"][0]["heading"] == "Principal Engineer | Analytical Engines"
    assert plan["resume_bullet_plan"][0]["rewrite_allowed"] is False
    assert plan["one_page_plan"]["target_pages"] == 1
    assert plan["cover_letter_required"] is True
    assert plan["cover_letter_evidence_points"]
    assert plan["project_substitution_suggestions"][0]["project"] == "Portal Control Center"


def test_detected_field_answer_uses_verified_profile_and_flags_sensitive_fields() -> None:
    canonical_profile = {
        "profile": {
            "legalName": "Ada Lovelace",
            "email": "ada@example.com",
            "workAuthorization": {"summary": "Requires explicit user confirmation"},
        }
    }

    email_answer = propose_answer_for_detected_field(
        field_label="Email address",
        field_type="text",
        canonical_profile=canonical_profile,
    )
    work_auth_answer = propose_answer_for_detected_field(
        field_label="Are you legally authorized to work?",
        field_type="radio",
        canonical_profile=canonical_profile,
    )

    assert email_answer.proposed_value == "ada@example.com"
    assert email_answer.requires_review is False
    assert work_auth_answer.requires_review is True
    assert work_auth_answer.source == "PROFILE"


def test_validator_blocks_banned_words_and_em_dash() -> None:
    report = TextArtifactValidator().validate("I utilize Python \u2014 often.", artifact_kind="cover_letter")

    assert report.passed is False
    assert {issue.code for issue in report.blocking_issues} == {"BANNED_WORD", "EM_DASH"}
    assert "Validation failed." in report.to_dict()["human_report"]


def test_validator_reads_docx_artifacts(tmp_path: Path) -> None:
    from docx import Document  # type: ignore

    source = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Experience")
    document.add_paragraph("Skills")
    document.add_paragraph("Built robust Python services.")
    document.save(str(source))

    report = TextArtifactValidator().validate_file(source, artifact_kind="resume")

    assert report.passed is False
    assert "BANNED_WORD" in {issue.code for issue in report.blocking_issues}


def test_validator_reads_tex_artifacts(tmp_path: Path) -> None:
    source = tmp_path / "cover.tex"
    source.write_text(
        r"""
\documentclass{article}
\begin{document}
\section*{Letter}
I utilize Python services.
\end{document}
""",
        encoding="utf-8",
    )

    report = TextArtifactValidator().validate_file(source, artifact_kind="cover_letter")

    assert report.passed is False
    assert "BANNED_WORD" in {issue.code for issue in report.blocking_issues}


def test_validator_warns_on_weak_resume_bullets() -> None:
    report = TextArtifactValidator().validate(
        "Experience\nSkills\n- Worked on internal tools\n- Worked on APIs\n- Worked on test suites",
        artifact_kind="resume",
    )

    assert report.passed is True
    assert "WEAK_BULLET" in {issue.code for issue in report.warnings}
    assert "REPETITIVE_RHYTHM" in {issue.code for issue in report.warnings}


def test_resume_artifact_generation_uses_verified_profile_only(tmp_path: Path) -> None:
    canonical_profile = {
        "profile": {"legalName": "Ada Lovelace", "email": "ada@example.com", "location": "London"},
        "skillGroups": [{"skills": ["Python", "SQLite", "TypeScript"]}],
        "experience": [{"title": "Principal Engineer", "company": "Analytical Engines", "bullets": ["Built deterministic data tools."]}],
        "projects": [],
        "education": [],
    }
    plan = {"matched_evidence": [{"keyword": "Python"}], "missing_keywords": ["Kubernetes"]}

    content = build_resume_markdown(canonical_profile=canonical_profile, tailoring_plan=plan)
    artifact = write_text_artifact(
        output_dir=tmp_path,
        filename_input=GeneratedNameInput(
            first_name="Ada",
            last_name="Lovelace",
            company="Example Co",
            role="Staff Engineer",
            kind="Resume",
            extension="md",
        ),
        content=content,
        file_kind="RESUME",
        extension="md",
        review_only=True,
    )

    assert "Kubernetes" in content
    assert "Missing keywords were not added" in content
    assert artifact.format == "MD"
    assert Path(artifact.local_path).exists()
    assert artifact.sha256


def test_resume_artifact_generation_applies_tailoring_plan_order_and_limits() -> None:
    canonical_profile = {
        "profile": {"legalName": "Ada Lovelace"},
        "skillGroups": [{"skills": ["TypeScript", "Python", "SQLite"]}],
        "experience": [
            {
                "id": "exp-old",
                "title": "Engineer",
                "company": "Older Co",
                "bullets": ["Maintained internal reporting tools."],
            },
            {
                "id": "exp-target",
                "title": "Principal Engineer",
                "company": "Analytical Engines",
                "bullets": [
                    "Built Python automation for SQLite workflows.",
                    "Reviewed TypeScript services.",
                    "Kept release diagnostics readable.",
                ],
            },
        ],
        "projects": [],
        "education": [],
    }
    plan = {
        "matched_evidence": [{"keyword": "Python"}, {"keyword": "SQLite"}],
        "missing_keywords": [],
        "resume_bullet_plan": [
            {
                "source_type": "experience",
                "source_id": "exp-target",
                "heading": "Principal Engineer | Analytical Engines",
                "keep_bullets": ["Built Python automation for SQLite workflows."],
            }
        ],
        "one_page_plan": {
            "experience_limit": 1,
            "project_limit": 0,
            "bullet_limit_per_role": 1,
            "skills_priority": ["Python", "SQLite"],
        },
    }

    content = build_resume_markdown(canonical_profile=canonical_profile, tailoring_plan=plan)

    assert "Principal Engineer | Analytical Engines" in content
    assert "Engineer | Older Co" not in content
    assert "Built Python automation for SQLite workflows." in content
    assert "Reviewed TypeScript services." not in content
    assert "Python, SQLite, TypeScript" in content


def test_cover_letter_artifact_refuses_missing_evidence() -> None:
    assert split_legal_name("Ada Lovelace") == ("Ada", "Lovelace")
    assert (
        build_cover_letter_text(
            canonical_profile={"profile": {"legalName": "Ada Lovelace"}, "experience": [], "projects": []},
            job_metadata={"company": "Example Co", "role": "Staff Engineer"},
            tailoring_plan={"matched_evidence": []},
        )
        is None
    )


def test_validator_warns_bullet_too_long_with_default_limit() -> None:
    long_bullet = "A" * 241
    report = TextArtifactValidator().validate(
        f"Experience\nSkills\n- {long_bullet}",
        artifact_kind="resume",
    )

    assert report.passed is True
    codes = {issue.code for issue in report.warnings}
    assert "BULLET_TOO_LONG" in codes


def test_validator_no_bullet_too_long_within_limit() -> None:
    bullet = "A" * 115
    report = TextArtifactValidator().validate(
        f"Experience\nSkills\n- {bullet}",
        artifact_kind="resume",
        bullet_char_limit=220,
    )

    codes = {issue.code for issue in report.warnings}
    assert "BULLET_TOO_LONG" not in codes


def test_validator_bullet_too_long_custom_limit() -> None:
    bullet = "A" * 221
    report = TextArtifactValidator().validate(
        f"Experience\nSkills\n- {bullet}",
        artifact_kind="resume",
        bullet_char_limit=220,
    )

    codes = {issue.code for issue in report.warnings}
    assert "BULLET_TOO_LONG" in codes
