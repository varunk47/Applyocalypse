from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path


def test_pipeline_validate_text_cli(tmp_path: Path) -> None:
    source = tmp_path / "cover.txt"
    source.write_text("I utilize Python", encoding="utf-8")

    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "applyocalypse_automation.runner",
            "pipeline",
            "validate-text",
            "--source-text-file",
            str(source),
            "--artifact-kind",
            "cover_letter",
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    report = json.loads(completed.stdout)
    assert report["passed"] is False
    assert report["blocking_issues"][0]["code"] == "BANNED_WORD"


def test_pipeline_validate_file_cli(tmp_path: Path) -> None:
    source = tmp_path / "resume.tex"
    source.write_text(
        r"\section*{Experience}\n\begin{itemize}\n\item Built seamless tooling.\n\end{itemize}",
        encoding="utf-8",
    )

    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "applyocalypse_automation.runner",
            "pipeline",
            "validate-file",
            "--source",
            str(source),
            "--artifact-kind",
            "resume",
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    report = json.loads(completed.stdout)
    assert report["passed"] is False
    assert report["blocking_issues"][0]["code"] == "BANNED_WORD"


def test_pipeline_repair_anchors_cli(tmp_path: Path) -> None:
    source = tmp_path / "resume.tex"
    output = tmp_path / "anchored.tex"
    source.write_text("\\section{Skills}\nPython, SQLite\n", encoding="utf-8")

    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "applyocalypse_automation.runner",
            "pipeline",
            "repair-anchors",
            "--source",
            str(source),
            "--output",
            str(output),
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    payload = json.loads(completed.stdout)
    assert payload["output_path"] == str(output)
    assert "{{APPLYO_SKILLS}}" in output.read_text(encoding="utf-8")


def test_pipeline_generated_name_cli(tmp_path: Path) -> None:
    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "applyocalypse_automation.runner",
            "pipeline",
            "generated-name",
            "--first-name",
            "Ada",
            "--last-name",
            "Lovelace",
            "--company",
            "Analytical Engines",
            "--role",
            "Staff Engineer",
            "--kind",
            "Resume",
            "--extension",
            "pdf",
            "--output-dir",
            str(tmp_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    payload = json.loads(completed.stdout)
    assert payload["filename"] == "Ada Lovelace Analytical Engines Staff Engineer Resume.pdf"


def test_pipeline_parse_source_cli_extracts_sections_and_skills(tmp_path: Path) -> None:
    source = tmp_path / "resume.txt"
    source.write_text(
        "\n".join(
            [
                "Ada Lovelace",
                "ada@example.com",
                "Skills",
                "Python, TypeScript, SQLite",
                "Experience",
                "- Built deterministic document tooling",
            ]
        ),
        encoding="utf-8",
    )

    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "applyocalypse_automation.runner",
            "pipeline",
            "parse-source",
            "--source",
            str(source),
            "--document-kind",
            "RESUME",
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    payload = json.loads(completed.stdout)
    assert payload["parser_name"] == "applyocalypse-local-parser"
    assert payload["canonical"]["identity"]["email"] == "ada@example.com"
    assert payload["canonical"]["sections"][0]["normalizedLabel"] == "skills"
    assert payload["canonical"]["skillGroups"][0]["skills"] == ["Python", "TypeScript", "SQLite"]


def test_pipeline_parse_source_cli_extracts_docx_resume(tmp_path: Path) -> None:
    from docx import Document

    source = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Ada Lovelace")
    document.add_paragraph("ada@example.com")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, TypeScript, SQLite")
    document.add_paragraph("Experience")
    document.add_paragraph("Built deterministic document tooling", style="List Bullet")
    document.save(source)

    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "applyocalypse_automation.runner",
            "pipeline",
            "parse-source",
            "--source",
            str(source),
            "--document-kind",
            "RESUME",
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    payload = json.loads(completed.stdout)
    assert payload["parser_name"] == "applyocalypse-local-parser"
    assert payload["canonical"]["identity"]["email"] == "ada@example.com"
    assert payload["canonical"]["skillGroups"][0]["skills"] == ["Python", "TypeScript", "SQLite"]
    assert payload["style_map"]["paragraphs"]


def test_worker_run_analyzes_local_job_text(tmp_path: Path) -> None:
    job_text = tmp_path / "job.txt"
    work_dir = tmp_path / "run"
    profile = tmp_path / "profile.json"
    job_text.write_text("Required: Python, TypeScript. Cover letter required.", encoding="utf-8")
    profile.write_text(
        json.dumps(
            {
                "profile": {
                    "legalName": "Ada Lovelace",
                    "email": "ada@example.com",
                    "phone": None,
                    "location": "London",
                    "workAuthorization": {},
                },
                "skillGroups": [{"label": "Languages", "skills": ["Python", "TypeScript"]}],
            }
        ),
        encoding="utf-8",
    )

    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "applyocalypse_automation.runner",
            "--run-id",
            "run-local",
            "--work-dir",
            str(work_dir),
            "--job-text-file",
            str(job_text),
            "--profile-json-file",
            str(profile),
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    events = [json.loads(line) for line in completed.stdout.splitlines()]
    event_types = [event["event_type"] for event in events]
    assert "JD_ANALYSIS_COMPLETED" in event_types
    assert "FIELD_VALUE_PROPOSED" in event_types
    assert (work_dir / "tailoring-plan.json").exists()
    assert events[-1]["event_type"] == "PAUSED"
