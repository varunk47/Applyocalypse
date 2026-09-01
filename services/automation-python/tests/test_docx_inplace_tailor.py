"""In-place resume tailoring must change bullet TEXT while preserving the master's
exact run formatting (font, size, bold) and paragraph list/bullet style."""
from __future__ import annotations

from docx import Document
from docx.shared import Pt

from applyocalypse_automation.documents import docx_mutation
from applyocalypse_automation.documents.docx_mutation import (
    addressable_paragraphs,
    collect_tailorable_bullets,
    parses_back_intact,
    tailor_master_docx_in_place,
)


def _make_master(path) -> None:
    doc = Document()
    doc.add_paragraph("EXPERIENCE")
    bullet = doc.add_paragraph(style="List Bullet")
    run = bullet.add_run("Built a data pipeline processing 10GB of logs using Python and Airflow daily")
    run.font.name = "Garamond"
    run.font.size = Pt(11)
    run.bold = True
    doc.save(str(path))


def test_collect_finds_the_bullet(tmp_path):
    master = tmp_path / "master.docx"
    _make_master(master)

    bullets = collect_tailorable_bullets(Document(str(master)))

    assert len(bullets) == 1
    _, text = bullets[0]
    assert "data pipeline" in text


def test_in_place_tailor_preserves_formatting(tmp_path):
    master = tmp_path / "master.docx"
    _make_master(master)
    index, _ = collect_tailorable_bullets(Document(str(master)))[0]
    new_text = "Engineered an ETL pipeline handling 10GB of logs with Python and Airflow"

    out = tmp_path / "out.docx"
    _, changed = tailor_master_docx_in_place(master, out, {index: new_text})

    assert changed == 1
    result = Document(str(out))
    paragraph = result.paragraphs[index]
    assert paragraph.text.strip() == new_text
    assert paragraph.style.name == "List Bullet"  # bullet/indent style preserved
    donor = paragraph.runs[0]
    assert donor.font.name == "Garamond"  # font preserved
    assert donor.font.size == Pt(11)  # size preserved
    assert donor.bold is True  # inline bold preserved


def test_no_changes_copies_master_verbatim(tmp_path):
    master = tmp_path / "m.docx"
    _make_master(master)
    out = tmp_path / "o.docx"

    _, changed = tailor_master_docx_in_place(master, out, {})

    assert changed == 0
    assert out.exists()
    # The bullet still has its original text and formatting.
    result = Document(str(out))
    index, text = collect_tailorable_bullets(result)[0]
    assert "data pipeline" in text
    assert result.paragraphs[index].runs[0].font.name == "Garamond"


def _make_table_master(path) -> None:
    """The two-column resume template: every word lives inside a table, so the document
    body has no bullets in it at all."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    cell = table.cell(0, 0)
    cell.text = "EXPERIENCE"
    bullet = cell.add_paragraph(style="List Bullet")
    run = bullet.add_run("Built a data pipeline processing 10GB of logs using Python and Airflow daily")
    run.font.name = "Garamond"
    run.font.size = Pt(11)
    doc.save(str(path))


def test_bullets_inside_a_table_are_found_and_tailored(tmp_path):
    """Reading only ``document.paragraphs`` finds nothing in a table-based resume, and
    nothing to change is indistinguishable from a resume that needed no changes: the run
    reports success and hands back the untailored master."""
    master = tmp_path / "table.docx"
    _make_table_master(master)

    assert [p.text for p in Document(str(master)).paragraphs if p.text.strip()] == []

    bullets = collect_tailorable_bullets(Document(str(master)))
    assert len(bullets) == 1
    index, text = bullets[0]
    assert "data pipeline" in text

    new_text = "Engineered an ETL pipeline handling 10GB of logs with Python and Airflow"
    out = tmp_path / "out.docx"
    _, changed = tailor_master_docx_in_place(master, out, {index: new_text})

    assert changed == 1
    written = addressable_paragraphs(Document(str(out)))[index]
    assert written.text.strip() == new_text
    assert written.style.name == "List Bullet"
    assert written.runs[0].font.name == "Garamond"


def test_a_merged_cell_is_addressed_once(tmp_path):
    """A cell merged across columns is returned once per column it spans. Addressing the
    same paragraph under two indices would make the read-back check see an index it never
    tailored holding tailored text, and throw the whole document away."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    header = table.cell(0, 0).merge(table.cell(0, 1))
    header_bullet = header.paragraphs[0]
    header_bullet.style = "List Bullet"
    header_bullet.add_run("Led the platform team that rebuilt log ingestion end to end")
    table.cell(1, 0).add_paragraph(style="List Bullet").add_run(
        "Built a data pipeline processing 10GB of logs using Python and Airflow daily"
    )
    master = tmp_path / "merged.docx"
    doc.save(str(master))

    bullets = collect_tailorable_bullets(Document(str(master)))
    assert sum(1 for _, text in bullets if "platform team" in text) == 1

    index = next(i for i, text in bullets if "platform team" in text)
    new_text = "Led the platform group that rebuilt log ingestion end to end"
    out = tmp_path / "out.docx"
    _, changed = tailor_master_docx_in_place(master, out, {index: new_text})

    assert changed == 1
    assert addressable_paragraphs(Document(str(out)))[index].text.strip() == new_text


def test_a_document_that_does_not_parse_back_falls_back_to_the_master(tmp_path, monkeypatch):
    """Redistributing text across a paragraph's runs is what preserves the formatting, and
    it is also where a character goes missing. When the file on disk does not read back as
    what was asked for, the user gets their own resume rather than a mangled one."""
    master = tmp_path / "master.docx"
    _make_master(master)
    index, original = collect_tailorable_bullets(Document(str(master)))[0]

    intact = docx_mutation.replace_paragraph_text_preserving_runs

    def drops_a_character(paragraph, replacement_text):
        intact(paragraph, replacement_text[:-1])

    monkeypatch.setattr(docx_mutation, "replace_paragraph_text_preserving_runs", drops_a_character)

    out = tmp_path / "out.docx"
    _, changed = tailor_master_docx_in_place(
        master, out, {index: "Engineered an ETL pipeline for 10GB of daily logs"}
    )

    assert changed == 0
    assert Document(str(out)).paragraphs[index].text.strip() == original


def test_parses_back_intact_reads_the_document_it_asked_for():
    assert parses_back_intact(["heading", "old bullet"], ["heading", "new bullet"], {1: "new bullet"})
    # A paragraph nobody tailored came back different.
    assert not parses_back_intact(["heading", "old"], ["heading!", "old"], {})
    # A tailored paragraph came back as something other than the text requested.
    assert not parses_back_intact(["heading", "old"], ["heading", "new"], {1: "newer"})
    # A paragraph went missing on the way to disk.
    assert not parses_back_intact(["heading", "old"], ["heading"], {})
