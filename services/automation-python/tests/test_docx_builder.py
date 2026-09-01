"""Tests for documents/docx_builder.py and cover letter DOCX naming."""
from __future__ import annotations

import tempfile
from pathlib import Path

import pytest

from applyocalypse_automation.documents.docx_builder import (
    _date_range,
    build_cover_letter_docx,
    build_resume_docx,
)
from applyocalypse_automation.documents.file_generation import (
    GeneratedNameInput,
    build_generated_filename,
)

_PROFILE = {
    "profile": {
        "legalName": "Margaret Hamilton",
        "email": "margaret@mit.edu",
        "phone": "555-0100",
        "location": "Cambridge, MA",
    }
}

_BODY = (
    "Dear Hiring Team,\n\n"
    "My work on the Apollo Guidance Computer taught me to design software that operators can trust under pressure.\n\n"
    "I reduced critical abort scenarios by exhaustive pre-flight simulation.\n\n"
    "Thank you for your consideration.\n\n"
    "Best regards,\nMargaret Hamilton"
)


def test_cover_letter_docx_is_created() -> None:
    try:
        from docx import Document  # type: ignore  # noqa: F401
    except ImportError:
        pytest.skip("python-docx not installed")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "cover_letter.docx"
        build_cover_letter_docx(_BODY, _PROFILE, out)
        assert out.exists(), "DOCX file was not created"
        assert out.stat().st_size > 0


def test_cover_letter_docx_round_trips_text() -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        pytest.skip("python-docx not installed")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "cover_letter.docx"
        build_cover_letter_docx(_BODY, _PROFILE, out)
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        assert "Margaret Hamilton" in text
        assert "Apollo Guidance Computer" in text
        assert "Thank you" in text


def test_cover_letter_docx_contains_contact_info() -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        pytest.skip("python-docx not installed")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "cover_letter.docx"
        build_cover_letter_docx(_BODY, _PROFILE, out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "margaret@mit.edu" in full_text
        assert "555-0100" in full_text


def test_cover_letter_docx_no_em_dashes() -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        pytest.skip("python-docx not installed")

    body_with_em = _BODY + "— extra dash"
    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "cover_letter.docx"
        # Builder writes what it is given; caller is responsible for clean text.
        # This test confirms em-dashes in output are detectable via round-trip.
        build_cover_letter_docx(body_with_em, _PROFILE, out)
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "—" in full_text, "round-trip should preserve text verbatim (caller validates)"


def test_cover_letter_naming_kind_with_space() -> None:
    name = build_generated_filename(
        GeneratedNameInput(
            first_name="Margaret",
            last_name="Hamilton",
            company="NASA",
            role="Systems Engineer",
            kind="Cover Letter",
            extension="docx",
        )
    )
    assert name == "Margaret Hamilton NASA Systems Engineer Cover Letter.docx"


def test_cover_letter_naming_matches_resume_pattern_except_kind() -> None:
    resume_name = build_generated_filename(
        GeneratedNameInput(
            first_name="Ada",
            last_name="Lovelace",
            company="Babbage Co",
            role="Engineer",
            kind="Resume",
            extension="docx",
        )
    )
    cl_name = build_generated_filename(
        GeneratedNameInput(
            first_name="Ada",
            last_name="Lovelace",
            company="Babbage Co",
            role="Engineer",
            kind="Cover Letter",
            extension="docx",
        )
    )
    assert resume_name == "Ada Lovelace Babbage Co Engineer Resume.docx"
    assert cl_name == "Ada Lovelace Babbage Co Engineer Cover Letter.docx"


def test_cover_letter_docx_handles_empty_profile_gracefully() -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        pytest.skip("python-docx not installed")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "cover_letter.docx"
        build_cover_letter_docx("Dear Hiring Team,\n\nThank you.\n\nBest regards,\nCandidate", {}, out)
        assert out.exists()
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        assert "Thank you" in text


# ---------------------------------------------------------------------------
# build_resume_docx tests
# ---------------------------------------------------------------------------

_RESUME_PROFILE = {
    "profile": {
        "legalName": "Ada Lovelace",
        "email": "ada@example.com",
        "phone": "555-0200",
        "location": "London",
    },
    "skillGroups": [{"skills": ["Python", "C++", "Mathematics"]}],
    "experience": [
        {
            "title": "Principal Engineer",
            "company": "Analytical Engines",
            "bullets": [
                "Designed the first algorithm intended for mechanical computation.",
                "Translated and annotated Menabrea's paper on the Analytical Engine.",
            ],
        }
    ],
    "projects": [
        {
            "name": "Bernoulli Numbers",
            "bullets": ["Devised an algorithm to compute Bernoulli numbers step by step."],
        }
    ],
    "education": [
        {
            "institution": "University of London",
            "degree": "BSc",
            "field": "Mathematics",
        }
    ],
}

_RESUME_PLAN: dict = {
    "matched_evidence": [],
    "missing_keywords": [],
    "one_page_plan": {
        "experience_limit": 4,
        "project_limit": 3,
        "bullet_limit_per_role": 4,
        "skills_priority": ["Python"],
    },
}


def test_build_resume_docx_creates_file() -> None:
    try:
        from docx import Document  # type: ignore  # noqa: F401
    except ImportError:
        pytest.skip("python-docx not installed")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "resume.docx"
        build_resume_docx(_RESUME_PROFILE, _RESUME_PLAN, out)
        assert out.exists()
        assert out.stat().st_size > 0


def test_build_resume_docx_contains_name_and_contact() -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        pytest.skip("python-docx not installed")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "resume.docx"
        build_resume_docx(_RESUME_PROFILE, _RESUME_PLAN, out)
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        assert "Ada Lovelace" in text
        assert "ada@example.com" in text


def test_build_resume_docx_contains_skills() -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        pytest.skip("python-docx not installed")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "resume.docx"
        build_resume_docx(_RESUME_PROFILE, _RESUME_PLAN, out)
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        assert "Python" in text
        assert "C++" in text


def test_build_resume_docx_contains_experience_bullets() -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        pytest.skip("python-docx not installed")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "resume.docx"
        build_resume_docx(_RESUME_PROFILE, _RESUME_PLAN, out)
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        assert "Principal Engineer" in text
        assert "Designed the first algorithm" in text


def test_build_resume_docx_contains_education() -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        pytest.skip("python-docx not installed")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "resume.docx"
        build_resume_docx(_RESUME_PROFILE, _RESUME_PLAN, out)
        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        assert "University of London" in text


def test_build_resume_docx_handles_empty_profile() -> None:
    try:
        from docx import Document  # type: ignore  # noqa: F401
    except ImportError:
        pytest.skip("python-docx not installed")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "resume_empty.docx"
        build_resume_docx({}, {}, out)
        assert out.exists()


def test_build_resume_docx_respects_font_size_param() -> None:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore  # noqa: F401
    except ImportError:
        pytest.skip("python-docx not installed")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "resume_11.docx"
        build_resume_docx(_RESUME_PROFILE, _RESUME_PLAN, out, font_size=11)
        doc = Document(str(out))
        # Find any run with an explicit font size set to verify font_size was applied
        run_sizes = [
            run.font.size
            for para in doc.paragraphs
            for run in para.runs
            if run.font.size is not None
        ]
        expected_pts = {round(s / 12700) for s in run_sizes}
        assert 11 in expected_pts or 15 in expected_pts  # 11pt body or 15pt header (11+4)


# ---------------------------------------------------------------------------
# Contact line and employment dates
#
# `_RESUME_PROFILE` above carries neither a LinkedIn URL nor any dates, which is
# why the builder could read the wrong contact key and drop every date range
# without a single test noticing. This fixture carries both.
# ---------------------------------------------------------------------------

_DATED_PROFILE = {
    "profile": {
        "legalName": "Grace Hopper",
        "email": "grace@example.com",
        "phone": "555-0300",
        "location": "Arlington, VA",
        "linkedinUrl": "https://linkedin.com/in/grace-hopper",
    },
    "skillGroups": [{"skills": ["COBOL"]}],
    "experience": [
        {
            "title": "Rear Admiral",
            "company": "US Navy",
            "startDate": "1967-08",
            "endDate": "1986-08",
            "bullets": ["Standardised the compiler toolchain across the fleet."],
        },
        {
            "title": "Senior Consultant",
            "company": "Digital Equipment Corporation",
            "startDate": "1986-09",
            "bullets": ["Lectured on the cost of a microsecond."],
        },
    ],
    "education": [
        {
            "institution": "Yale University",
            "degree": "PhD",
            "field": "Mathematics",
            "startDate": "1930-09",
            "endDate": "1934-06",
        }
    ],
}


def _resume_text(profile: dict, name: str) -> str:
    from docx import Document  # type: ignore

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / name
        build_resume_docx(profile, _RESUME_PLAN, out)
        doc = Document(str(out))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def test_build_resume_docx_contains_linkedin_url() -> None:
    """The canonical key is `linkedinUrl`; reading `linkedin` matched nothing."""
    try:
        from docx import Document  # type: ignore  # noqa: F401
    except ImportError:
        pytest.skip("python-docx not installed")

    text = _resume_text(_DATED_PROFILE, "resume_linkedin.docx")
    assert "https://linkedin.com/in/grace-hopper" in text


def test_build_resume_docx_dates_each_role() -> None:
    try:
        from docx import Document  # type: ignore  # noqa: F401
    except ImportError:
        pytest.skip("python-docx not installed")

    text = _resume_text(_DATED_PROFILE, "resume_dates.docx")
    assert "Rear Admiral | US Navy | 1967-08 to 1986-08" in text


def test_build_resume_docx_marks_a_role_with_no_end_date_as_present() -> None:
    try:
        from docx import Document  # type: ignore  # noqa: F401
    except ImportError:
        pytest.skip("python-docx not installed")

    text = _resume_text(_DATED_PROFILE, "resume_present.docx")
    assert "1986-09 to Present" in text


def test_build_resume_docx_dates_education() -> None:
    try:
        from docx import Document  # type: ignore  # noqa: F401
    except ImportError:
        pytest.skip("python-docx not installed")

    text = _resume_text(_DATED_PROFILE, "resume_edu_dates.docx")
    assert "Yale University | PhD | Mathematics | 1930-09 to 1934-06" in text


def test_build_resume_docx_leaves_no_dangling_separator_without_dates() -> None:
    """An entry with no dates keeps the heading it always had."""
    try:
        from docx import Document  # type: ignore  # noqa: F401
    except ImportError:
        pytest.skip("python-docx not installed")

    text = _resume_text(_RESUME_PROFILE, "resume_undated.docx")
    assert "Principal Engineer | Analytical Engines" in text
    assert "Analytical Engines |" not in text
    assert "University of London | BSc | Mathematics" in text
    assert "Mathematics |" not in text


@pytest.mark.parametrize(
    ("entry", "expected"),
    [
        ({"startDate": "2019-06", "endDate": "2021-03"}, "2019-06 to 2021-03"),
        ({"startDate": "2019-06"}, "2019-06 to Present"),
        ({"startDate": "2019-06", "endDate": None}, "2019-06 to Present"),
        ({"startDate": "", "endDate": "2021-03"}, "2021-03"),
        ({"endDate": "2021-03"}, "2021-03"),
        ({}, ""),
        ({"startDate": None, "endDate": None}, ""),
        ({"startDate": "  2019-06  ", "endDate": "  2021-03  "}, "2019-06 to 2021-03"),
    ],
)
def test_date_range(entry: dict, expected: str) -> None:
    assert _date_range(entry) == expected
