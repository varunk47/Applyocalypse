"""Tests for font size detection and PDF page counting."""
from __future__ import annotations

from pathlib import Path

import pytest

from applyocalypse_automation.documents.font_detection import (
    _FALLBACK_FONT_SIZE,
    count_pdf_pages,
    detect_docx_body_font_size,
    detect_resume_font_size,
    detect_tex_body_font_size,
)


# ---------------------------------------------------------------------------
# detect_tex_body_font_size
# ---------------------------------------------------------------------------


def test_tex_detects_11pt_from_documentclass(tmp_path: Path) -> None:
    tex = tmp_path / "resume.tex"
    tex.write_text(r"\documentclass[11pt,letterpaper]{article}\begin{document}text\end{document}", encoding="utf-8")
    assert detect_tex_body_font_size(tex) == 11


def test_tex_detects_10pt_from_documentclass(tmp_path: Path) -> None:
    tex = tmp_path / "resume.tex"
    tex.write_text(r"\documentclass[10pt]{article}\begin{document}text\end{document}", encoding="utf-8")
    assert detect_tex_body_font_size(tex) == 10


def test_tex_falls_back_when_no_size_option(tmp_path: Path) -> None:
    tex = tmp_path / "resume.tex"
    tex.write_text(r"\documentclass{article}\begin{document}text\end{document}", encoding="utf-8")
    assert detect_tex_body_font_size(tex) == _FALLBACK_FONT_SIZE


def test_tex_detects_size_from_fontsize_command(tmp_path: Path) -> None:
    tex = tmp_path / "resume.tex"
    tex.write_text(r"\documentclass{article}\fontsize{11}{13}\begin{document}hi\end{document}", encoding="utf-8")
    assert detect_tex_body_font_size(tex) == 11


def test_tex_falls_back_for_unsupported_size(tmp_path: Path) -> None:
    tex = tmp_path / "resume.tex"
    tex.write_text(r"\documentclass[12pt]{article}\begin{document}text\end{document}", encoding="utf-8")
    assert detect_tex_body_font_size(tex) == _FALLBACK_FONT_SIZE


def test_tex_returns_fallback_on_missing_file() -> None:
    assert detect_tex_body_font_size(Path("/nonexistent/resume.tex")) == _FALLBACK_FONT_SIZE


# ---------------------------------------------------------------------------
# detect_docx_body_font_size
# ---------------------------------------------------------------------------


def test_docx_detects_11pt(tmp_path: Path) -> None:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except ImportError:
        pytest.skip("python-docx not installed")

    path = tmp_path / "resume.docx"
    doc = Document()
    for _ in range(5):
        para = doc.add_paragraph()
        run = para.add_run("Some bullet text here.")
        run.font.size = Pt(11)
    doc.save(str(path))

    assert detect_docx_body_font_size(path) == 11


def test_docx_detects_10pt(tmp_path: Path) -> None:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except ImportError:
        pytest.skip("python-docx not installed")

    path = tmp_path / "resume.docx"
    doc = Document()
    for _ in range(5):
        para = doc.add_paragraph()
        run = para.add_run("Some bullet text here.")
        run.font.size = Pt(10)
    doc.save(str(path))

    assert detect_docx_body_font_size(path) == 10


def test_docx_falls_back_when_no_font_size_set(tmp_path: Path) -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        pytest.skip("python-docx not installed")

    path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("No explicit font size.")
    doc.save(str(path))

    assert detect_docx_body_font_size(path) == _FALLBACK_FONT_SIZE


def test_docx_returns_fallback_on_missing_file() -> None:
    assert detect_docx_body_font_size(Path("/nonexistent/resume.docx")) == _FALLBACK_FONT_SIZE


# ---------------------------------------------------------------------------
# detect_resume_font_size dispatcher
# ---------------------------------------------------------------------------


def test_detect_resume_font_size_dispatches_to_tex(tmp_path: Path) -> None:
    tex = tmp_path / "r.tex"
    tex.write_text(r"\documentclass[11pt]{article}\begin{document}hi\end{document}", encoding="utf-8")
    assert detect_resume_font_size(tex) == 11


def test_detect_resume_font_size_returns_fallback_for_pdf() -> None:
    assert detect_resume_font_size(Path("resume.pdf")) == _FALLBACK_FONT_SIZE


# ---------------------------------------------------------------------------
# count_pdf_pages
# ---------------------------------------------------------------------------


def test_count_pdf_pages_returns_zero_on_missing_file() -> None:
    assert count_pdf_pages(Path("/nonexistent/file.pdf")) == 0
